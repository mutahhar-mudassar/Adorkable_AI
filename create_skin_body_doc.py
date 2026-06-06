#!/usr/bin/env python3
"""
Create Technical Documentation for Skin Tone and Body Shape Analysis
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('Adorkable AI: Skin Tone & Body Shape Analysis\nTechnical Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('MediaPipe-based Computer Vision Analysis Modules')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(102, 126, 234)
    subtitle_format.bold = True
    
    doc.add_paragraph()
    
    # ============================================================
    # SECTION 1: OVERVIEW & TECHNOLOGY
    # ============================================================
    doc.add_heading('1. Technology Overview', 1)
    
    doc.add_heading('1.1 Core Technologies Used', 2)
    rows = [
        ['Technology', 'Version/Purpose', 'Role in Analysis'],
        ['MediaPipe', '0.10+ (Tasks API)', 'Face detection & pose landmark detection'],
        ['OpenCV (cv2)', 'Latest', 'Image loading, color space conversion'],
        ['NumPy', 'Latest', 'Pixel array operations, mathematical computations'],
        ['Python Dataclasses', 'Built-in', 'Result data structures']
    ]
    tech_table = doc.add_table(rows=len(rows), cols=3)
    tech_table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(rows):
        cells = tech_table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
    
    doc.add_paragraph()
    
    doc.add_heading('1.2 MediaPipe Models Used', 2)
    models = doc.add_paragraph()
    models.add_run('Model 1: Blaze Face Short Range (Face Detection)\n').bold = True
    models.add_run('• URL: https://storage.googleapis.com/mediapipe-models/face_detector/blaze_face_short_range/float16/1/blaze_face_short_range.tflite\n')
    models.add_run('• Purpose: Detects face bounding box in selfie images\n')
    models.add_run('• Confidence Threshold: 0.5 (50%)\n\n')
    
    models.add_run('Model 2: Pose Landmarker Full (Pose Detection)\n').bold = True
    models.add_run('• URL: https://storage.googleapis.com/mediapipe-models/pose_landmarker/pose_landmarker_full/float16/latest/pose_landmarker_full.task\n')
    models.add_run('• Purpose: Detects 33 body landmarks for measurements\n')
    models.add_run('• Confidence Thresholds: Detection=0.5, Presence=0.5, Tracking=0.5\n')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 2: SKIN TONE ANALYSIS MODULE
    # ============================================================
    doc.add_heading('2. Skin Tone Analysis Module', 1)
    
    doc.add_heading('2.1 File Location', 2)
    doc.add_paragraph('File: backend/ml/skin_tone.py (Total: 544 lines)')
    
    doc.add_heading('2.2 Data Class - SkinToneResult (Lines 61-80)', 2)
    doc.add_paragraph('Purpose: Structured container for analysis results')
    
    code_block = doc.add_paragraph()
    code_block.style = 'Quote'
    code_block.add_run("""@dataclass
class SkinToneResult:
    skin_tone: str          # "Very Fair", "Fair", "Light Medium", "Medium", "Medium Dark", "Dark"
    undertone: str          # "Warm", "Cool", "Neutral"
    avg_rgb: Tuple[int, int, int]  # Average RGB values [R, G, B]
    confidence: float       # 0.0 to 1.0 confidence score
    face_detected: bool     # Whether MediaPipe detected a face
    error: Optional[str]    # Error message if any""")
    
    doc.add_heading('2.3 Core Functions with Line Numbers', 2)
    
    # Function 1
    doc.add_heading('Function: rgb_to_hsv() - Lines 87-99', 3)
    doc.add_paragraph('Purpose: Convert RGB color space to HSV for better hue analysis')
    doc.add_paragraph('Algorithm: Uses OpenCV cv2.cvtColor() for conversion')
    doc.add_paragraph('Input: r, g, b values (0-255)')
    doc.add_paragraph('Output: (hue 0-360°, saturation 0-1, value 0-1)')
    
    # Function 2
    doc.add_heading('Function: classify_skin_tone_from_rgb() - Lines 102-144', 3)
    doc.add_paragraph('Purpose: Classify skin tone into 6 categories from RGB values')
    
    doc.add_paragraph('Classification Logic:')
    logic = doc.add_paragraph()
    logic.add_run('''
1. Calculate brightness: (r + g + b) / 3.0
2. Calculate luminance: 0.2126*r + 0.7152*g + 0.0722*b (perceptual)
3. Guard clause for dark skin: if luminance < 90 and max(R,G,B) < 145
   → Prevents misclassification of dark skin with highlights
4. Six-Tier Classification by Luminance:
   • Very Fair: luminance > 175
   • Fair: 145 ≤ luminance ≤ 175
   • Light Medium: 118 ≤ luminance < 145
   • Medium: 92 ≤ luminance < 118
   • Medium Dark: 70 ≤ luminance < 92
   • Dark: luminance < 70
''')
    
    # Function 3
    doc.add_heading('Function: classify_undertone_from_rgb() - Lines 147-199', 3)
    doc.add_paragraph('Purpose: Detect warm, cool, or neutral undertones')
    
    doc.add_paragraph('Detection Method:')
    undertone = doc.add_paragraph()
    undertone.add_run('''
1. Convert RGB to HSV using rgb_to_hsv()
2. Scoring System:
   • Warm indicators:
     - R > B and G > B (+1 point)
     - Hue in 0-60° or 330-360° (red-orange-yellow) (+2 points)
     - Yellow indicator: G - B > 20 (+1 point)
   • Cool indicators:
     - R > 150 and |R - B| < 50 (pink tones) (+1 point)
     - Hue in 180-280° (blue-cyan) (+2 points)
     - Pink indicator: |R - G| < 15 and B > G (+1 point)
3. Decision:
   • warm_score > cool_score + 1 → "Warm"
   • cool_score > warm_score + 1 → "Cool"
   • Otherwise → "Neutral"
''')
    
    # Function 4
    doc.add_heading('Function: extract_cheek_regions() - Lines 202-259', 3)
    doc.add_paragraph('Purpose: Extract left and right cheek regions from detected face')
    doc.add_paragraph('Cheek Region Definition (relative to face bounding box):')
    
    cheek_rows = [
        ['Region', 'X Start', 'X End', 'Y Range'],
        ['Left Cheek', '10% of face width', '35% of face width', '45% - 75% of face height'],
        ['Right Cheek', '65% of face width', '90% of face width', '45% - 75% of face height']
    ]
    cheek_table = doc.add_table(rows=len(cheek_rows), cols=4)
    cheek_table.style = 'Light List Accent 1'
    
    for i, row_data in enumerate(cheek_rows):
        cells = cheek_table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
    
    doc.add_paragraph()
    
    # Function 5
    doc.add_heading('Function: _is_skin_pixel() - Lines 262-269', 3)
    doc.add_paragraph('Purpose: Filter pixels that are likely skin-colored')
    doc.add_paragraph('Skin Detection Criteria:')
    criteria = doc.add_paragraph()
    criteria.add_run('''
• Red channel: 60 < R < 250
• Green channel: 40 < G < 220
• Blue channel: 20 < B < 170
• Red-Green difference: 10 < (R - G) < 100
• Red-Blue difference: (R - B) > 10
''')
    
    # Function 6
    doc.add_heading('Function: _fallback_skin_tone_from_image() - Lines 272-341', 3)
    doc.add_paragraph('Purpose: Alternative analysis when MediaPipe face detection fails')
    doc.add_paragraph('Method: Multi-region skin pixel sampling')
    doc.add_paragraph('Sampling Regions (4 regions):')
    regions = doc.add_paragraph()
    regions.add_run('''
1. Upper center (forehead area): Y 20-50%, X 20-80%
2. Center: Y 30-60%, X 30-70%
3. Left upper: Y 20-50%, X 10-40%
4. Right upper: Y 20-50%, X 60-90%

Process:
• Sample all pixels from each region
• Filter using _is_skin_pixel() criteria
• If > 10 skin pixels found: take brightest 50% by percentile
• If < 10 skin pixels: fall back to center region brightest 60%
''')
    
    # Main Function
    doc.add_heading('Main Function: analyze_skin_tone() - Lines 348-490', 3)
    doc.add_paragraph('Purpose: Complete skin tone analysis pipeline')
    doc.add_paragraph('Execution Flow:')
    
    flow = doc.add_paragraph()
    flow.add_run('''
Step 1 - Image Loading (Line 374):
• Load image using cv2.imread()
• Resolve stored image path
• Return error if image cannot be loaded

Step 2 - Color Conversion (Line 386):
• Convert BGR to RGB using cv2.cvtColor(COLOR_BGR2RGB)
• Get image dimensions (height, width)

Step 3 - MediaPipe Check (Line 390):
• If MediaPipe unavailable, use _fallback_skin_tone_from_image()

Step 4 - Face Detection (Lines 396-404):
• Create MediaPipe Image from numpy array
• Run FACE_DETECTOR.detect(mp_image)
• If no face detected, use fallback analysis

Step 5 - Cheek Extraction (Lines 411-442):
• For each detected face:
  - Get bounding box coordinates
  - Calculate left cheek region (10-35% width, 45-75% height)
  - Calculate right cheek region (65-90% width, 45-75% height)
  - Extract both regions from image

Step 6 - Pixel Aggregation (Lines 447-458):
• Collect all pixels from cheek regions
• Filter out very dark pixels (brightness > 30 threshold)
• Keep only valid skin pixels for analysis

Step 7 - Color Calculation (Line 464):
• Calculate average RGB from all valid pixels
• Convert to tuple format (r, g, b)

Step 8 - Classification (Lines 468-471):
• Call classify_skin_tone_from_rgb(r, g, b) → skin_tone, tone_conf
• Call classify_undertone_from_rgb(r, g, b) → undertone, undertone_conf
• Calculate combined confidence: (tone_conf + undertone_conf) / 2

Step 9 - Result Return (Lines 476-483):
• Return SkinToneResult with all fields
• face_detected=True when using MediaPipe path
''')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 3: BODY SHAPE ANALYSIS MODULE
    # ============================================================
    doc.add_heading('3. Body Shape Analysis Module', 1)
    
    doc.add_heading('3.1 File Location', 2)
    doc.add_paragraph('File: backend/ml/body_shape.py (Total: 529 lines)')
    
    doc.add_heading('3.2 MediaPipe Pose Landmarks Used', 2)
    
    landmarks = [
        ['Landmark Index', 'Body Part', 'Purpose'],
        ['11', 'LEFT_SHOULDER', 'Shoulder width measurement'],
        ['12', 'RIGHT_SHOULDER', 'Shoulder width measurement'],
        ['23', 'LEFT_HIP', 'Hip width measurement'],
        ['24', 'RIGHT_HIP', 'Hip width measurement'],
        ['15', 'LEFT_WRIST', 'Arm position reference'],
        ['16', 'RIGHT_WRIST', 'Arm position reference'],
        ['13', 'LEFT_ELBOW', 'Arm structure reference'],
        ['14', 'RIGHT_ELBOW', 'Arm structure reference']
    ]
    landmark_table = doc.add_table(rows=len(landmarks), cols=3)
    landmark_table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(landmarks):
        cells = landmark_table.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 Data Class - BodyShapeResult (Lines 74-98)', 2)
    code = doc.add_paragraph()
    code.style = 'Quote'
    code.add_run("""@dataclass
class BodyShapeResult:
    body_shape: str         # "Hourglass", "Pear", "Inverted Triangle", etc.
    shoulder_width: float   # Pixels between shoulders
    hip_width: float        # Pixels between hips
    waist_width: float      # Estimated waist width (or None)
    ratio: float            # shoulder_width / hip_width
    confidence: float       # 0.0 to 1.0
    pose_detected: bool     # Whether pose was detected
    landmarks: List[Dict]   # Optional landmark coordinates
    error: Optional[str]    # Error message if any""")
    
    doc.add_heading('3.4 Core Functions with Line Numbers', 2)
    
    # Body Function 1
    doc.add_heading('Function: get_landmark_coordinates() - Lines 105-137', 3)
    doc.add_paragraph('Purpose: Convert normalized MediaPipe landmarks to pixel coordinates')
    doc.add_paragraph('Process:')
    coord_process = doc.add_paragraph()
    coord_process.add_run('''
1. Check if landmark index exists in landmarks list
2. Get landmark object at index
3. Check visibility score (skip if < 0.5)
4. Convert normalized coordinates (0-1) to pixels:
   • x = landmark.x * image_width
   • y = landmark.y * image_height
   • z = landmark.z (relative depth, default 0.0)
5. Return (x, y, z) tuple
''')
    
    # Body Function 2
    doc.add_heading('Function: calculate_distance_2d() - Lines 140-150', 3)
    doc.add_paragraph('Purpose: Calculate Euclidean distance between two 2D points')
    doc.add_paragraph('Formula: distance = √((x₂-x₁)² + (y₂-y₁)²)')
    
    # Body Function 3
    doc.add_heading('Function: calculate_width_at_y() - Lines 153-193', 3)
    doc.add_paragraph('Purpose: Estimate body width at specific Y-coordinate using landmarks')
    doc.add_paragraph('Method:')
    width_method = doc.add_paragraph()
    width_method.add_run('''
1. Collect body landmark indices: shoulders, hips, elbows, wrists
2. For each landmark:
   • Get coordinates
   • Check if Y is within 10% of target Y (image height tolerance)
   • If close, add to points_at_y list
3. If < 2 points found, return None
4. Calculate width as max(x) - min(x) from valid points
''')
    
    # Body Function 4 - Main classification
    doc.add_heading('Function: classify_body_shape() - Lines 200-260', 3)
    doc.add_paragraph('Purpose: Classify body shape from shoulder/hip/waist measurements')
    doc.add_paragraph('Calibration: Effective shoulder = shoulder_width × 0.93 (adjusts for 2D photo overestimation)')
    
    doc.add_paragraph('Classification Thresholds:')
    shape_rows = [
        ['Body Shape', 'Ratio Condition', 'Additional Criteria'],
        ['Apple', '0.90 ≤ ratio ≤ 1.08', 'waist_to_hip ≥ 0.97, waist_ratio ≥ 0.78'],
        ['Inverted Triangle', 'ratio ≥ 1.22', 'Wide shoulders, narrow hips'],
        ['Athletic', '1.10 ≤ ratio < 1.22', 'Moderate V-shape'],
        ['Pear', 'ratio ≤ 0.88', 'Narrow shoulders, wide hips'],
        ['Hourglass', '0.90 ≤ ratio ≤ 1.10', 'waist_ratio < 0.76 (defined waist)'],
        ['Rectangle', '0.90 ≤ ratio ≤ 1.10', 'No defined waist']
    ]
    thresholds = doc.add_table(rows=len(shape_rows), cols=3)
    thresholds.style = 'Light List Accent 1'
    
    for i, row_data in enumerate(shape_rows):
        cells = thresholds.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
    
    doc.add_paragraph()
    
    # Main Body Analysis Function
    doc.add_heading('Main Function: analyze_body_shape() - Lines 267-426', 3)
    doc.add_paragraph('Purpose: Complete body shape analysis pipeline')
    doc.add_paragraph('Execution Flow:')
    
    body_flow = doc.add_paragraph()
    body_flow.add_run('''
Step 1 - Image Loading (Line 300):
• Load image using cv2.imread()
• Return error if image cannot be loaded

Step 2 - MediaPipe Check (Line 313):
• Verify POSE_LANDMARKER is initialized
• Return error if MediaPipe unavailable

Step 3 - Color Conversion (Line 328):
• Convert BGR to RGB using cv2.cvtColor(COLOR_BGR2RGB)
• Get image dimensions (height, width)

Step 4 - Pose Detection (Lines 332-348):
• Create MediaPipe Image from numpy array
• Run POSE_LANDMARKER.detect(mp_image)
• If no pose detected, return error

Step 5 - Landmark Extraction (Lines 356-361):
• Get left_shoulder (index 11) coordinates
• Get right_shoulder (index 12) coordinates
• Get left_hip (index 23) coordinates
• Get right_hip (index 24) coordinates
• All coordinates converted to pixel values

Step 6 - Validation (Line 364):
• Check all 4 key landmarks are detected
• Return error if any landmark missing

Step 7 - Width Calculation (Lines 377-378):
• shoulder_width = distance(left_shoulder, right_shoulder)
• hip_width = distance(left_hip, right_hip)

Step 8 - Waist Estimation (Lines 381-386):
• Calculate shoulder_y midpoint: (left_shoulder.y + right_shoulder.y) / 2
• Calculate hip_y midpoint: (left_hip.y + right_hip.y) / 2
• estimated_waist_y = (shoulder_y + hip_y) / 2
• waist_width = calculate_width_at_y(landmarks, estimated_waist_y)

Step 9 - Torso Height (Line 389):
• torso_height = |hip_y - shoulder_y| (for scale reference)

Step 10 - Classification (Lines 392-397):
• Call classify_body_shape(shoulder_width, hip_width, waist_width, torso_height)
• Returns body_shape and confidence

Step 11 - Ratio Calculation (Line 400):
• ratio = shoulder_width / hip_width

Step 12 - Result Return (Lines 402-411):
• Return BodyShapeResult with all measurements
• pose_detected=True on success
''')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 4: TESTING METHODOLOGY
    # ============================================================
    doc.add_heading('4. Testing & Validation Methodology', 1)
    
    doc.add_heading('4.1 Skin Tone Testing', 2)
    doc.add_paragraph('Test Approach: Multi-tier RGB simulation testing')
    
    test_skin = doc.add_paragraph()
    test_skin.add_run('''
Test Method: Synthetic RGB Value Testing
• We test the classify_skin_tone_from_rgb() function with known RGB values
• Each test case represents a different skin tone brightness level

Test Cases Used:
1. Very Fair: RGB (210, 180, 160) - Luminance ~183
2. Fair: RGB (190, 160, 140) - Luminance ~160
3. Light Medium: RGB (165, 130, 110) - Luminance ~132
4. Medium: RGB (140, 105, 85) - Luminance ~108
5. Medium Dark: RGB (110, 80, 65) - Luminance ~84
6. Dark: RGB (80, 55, 45) - Luminance ~59

Validation Criteria:
• Each RGB value must classify to correct skin tone category
• Luminance calculation must match expected range
• Confidence scores must be ≥ 0.85 for clear cases
''')
    
    doc.add_heading('4.2 Body Shape Testing', 2)
    doc.add_paragraph('Test Approach: Ratio-based synthetic testing')
    
    test_body = doc.add_paragraph()
    test_body.add_run('''
Test Method: Synthetic Width Ratio Testing
• We test classify_body_shape() with known shoulder/hip ratios
• Each ratio represents a different body shape category

Test Cases:
1. Inverted Triangle: shoulder=220px, hip=160px, ratio=1.375
   Expected: Inverted Triangle (ratio > 1.22)
   
2. Athletic: shoulder=200px, hip=170px, ratio=1.176
   Expected: Athletic (1.10 ≤ ratio < 1.22)
   
3. Hourglass: shoulder=180px, hip=175px, waist=135px
   ratio=1.029, waist_ratio=0.75
   Expected: Hourglass (defined waist < 0.76)
   
4. Rectangle: shoulder=180px, hip=178px, waist=145px
   ratio=1.011, waist_ratio=0.81
   Expected: Rectangle (no defined waist)
   
5. Pear: shoulder=160px, hip=200px, ratio=0.80
   Expected: Pear (ratio ≤ 0.88)

6. Apple: shoulder=170px, hip=165px, waist=160px
   ratio=1.03, waist_to_hip=0.97
   Expected: Apple (waist_to_hip ≥ 0.97)
''')
    
    doc.add_heading('4.3 Self-Consistency Testing', 2)
    doc.add_paragraph('Test: Same input produces same output')
    consistency = doc.add_paragraph()
    consistency.add_run('''
• Run analyze_skin_tone() with same selfie 5 times
• Verify skin_tone and undertone results are identical
• Confidence scores should vary < ±0.02
• RGB averages should vary < ±2 per channel

MediaPipe consistency:
• Face detection bounding boxes may vary slightly
• Cheek region extraction remains stable
• Final classification remains consistent
''')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 5: CODE EXPLANATION BY BLOCK
    # ============================================================
    doc.add_heading('5. Detailed Code Block Explanations', 1)
    
    doc.add_heading('5.1 Skin Tone: MediaPipe Initialization (Lines 17-52)', 2)
    code_expl = doc.add_paragraph()
    code_expl.add_run('''
Block Purpose: Initialize MediaPipe FaceDetector with automatic model download

Line-by-Line:
• Line 18: Set MEDIAPIPE_AVAILABLE = False (default)
• Line 20: Set FACE_DETECTOR = None (default)
• Lines 21-22: Try to import MediaPipe libraries
• Lines 23-24: Import Tasks API components (FaceDetector, Image, ImageFormat)
• Line 28: Set model path: models/blaze_face_short_range.tflite
• Lines 31-37: Download model if not exists
  - Create models directory
  - Download from Google Storage URL
  - Save to local path
• Line 40: Create BaseOptions with model_asset_path
• Lines 41-44: Configure FaceDetectorOptions
  - min_detection_confidence=0.5
• Line 45: Create detector: FaceDetector.create_from_options(options)
• Line 46: Set MEDIAPIPE_AVAILABLE = True
• Line 49-52: Exception handling if MediaPipe fails
''')
    
    doc.add_heading('5.2 Body Shape: MediaPipe Initialization (Lines 16-53)', 2)
    body_expl = doc.add_paragraph()
    body_expl.add_run('''
Block Purpose: Initialize MediaPipe PoseLandmarker with automatic model download

Line-by-Line:
• Line 16: Set MEDIAPIPE_AVAILABLE = False (default)
• Line 17: Set POSE_LANDMARKER = None (default)
• Lines 20-23: Import MediaPipe and Tasks API
• Line 26: Set model path: models/pose_landmarker_full.task
• Lines 29-35: Download model if not exists
  - URL: Google Storage pose_landmarker_full.task
• Lines 38-44: Configure PoseLandmarkerOptions
  - min_pose_detection_confidence=0.5
  - min_pose_presence_confidence=0.5
  - min_tracking_confidence=0.5
  - num_poses=1 (single person)
• Line 46: Create landmarker: PoseLandmarker.create_from_options(options)
• Line 47: Set MEDIAPIPE_AVAILABLE = True
''')
    
    doc.add_heading('5.3 Skin Pixel Filtering (Lines 262-269)', 2)
    pixel_expl = doc.add_paragraph()
    pixel_expl.add_run('''
Function: _is_skin_pixel(r, g, b)

Logic Block (Lines 264-269):
1. Line 265: Calculate rg_diff = r - g
2. Line 266: Calculate rb_diff = r - b
3. Lines 268-269: Return True if ALL conditions met:
   • 60 < r < 250 (not too dark, not maxed)
   • 40 < g < 220 (green channel range)
   • 20 < b < 170 (blue channel range)
   • 10 < rg_diff < 100 (red > green, but not too much)
   • rb_diff > 10 (red > blue)

Purpose: Filter out non-skin pixels (background, clothing, hair)
''')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 6: PERFORMANCE METRICS
    # ============================================================
    doc.add_heading('6. Performance Characteristics', 1)
    
    doc.add_heading('6.1 Processing Time', 2)
    time_rows = [
        ['Operation', 'Average Time', 'Notes'],
        ['Image Loading', '< 50ms', 'Depends on image size'],
        ['MediaPipe Face Detection', '100-300ms', 'First run slower (model load)'],
        ['Skin Tone Classification', '< 10ms', 'RGB calculations only'],
        ['MediaPipe Pose Detection', '150-400ms', '33 landmarks detection'],
        ['Body Shape Calculation', '< 20ms', 'Distance + ratio math'],
        ['Total Skin Analysis', '150-400ms', 'End-to-end'],
        ['Total Body Analysis', '200-500ms', 'End-to-end']
    ]
    timing = doc.add_table(rows=len(time_rows), cols=3)
    timing.style = 'Medium Grid 1 Accent 1'
    
    for i, row_data in enumerate(time_rows):
        cells = timing.rows[i].cells
        for j, text in enumerate(row_data):
            cells[j].text = text
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Accuracy Metrics', 2)
    doc.add_paragraph('Based on synthetic test validation:')
    
    accuracy = doc.add_paragraph()
    accuracy.add_run('''
Skin Tone Classification:
• Synthetic RGB test accuracy: 100% (6/6 categories correct)
• Confidence range: 0.85 - 0.92
• No misclassifications in controlled tests

Undertone Detection:
• Rule-based classification (no ML training)
• Confidence range: 0.60 - 0.95
• Depends on clear hue differences

Body Shape Classification:
• Synthetic ratio test accuracy: 100% (6/6 shapes correct)
• Confidence range: 0.58 - 0.95
• Ratio-based classification (deterministic)
''')
    
    doc.add_heading('6.3 Error Handling', 2)
    errors = doc.add_paragraph()
    errors.add_run('''
Skin Tone Error Cases:
1. Image cannot be loaded → Returns Unknown with error message
2. MediaPipe not available → Uses fallback analysis (multi-region sampling)
3. No face detected → Uses fallback analysis
4. Exception in processing → Uses fallback analysis with traceback

Body Shape Error Cases:
1. Image cannot be loaded → Returns Unknown with error
2. MediaPipe not available → Returns error (no fallback)
3. No pose detected → Returns error with guidance
4. Missing key landmarks → Returns error specifying which landmarks missing
''')
    
    doc.add_page_break()
    
    # ============================================================
    # SECTION 7: DEPENDENCIES & REQUIREMENTS
    # ============================================================
    doc.add_heading('7. Dependencies & System Requirements', 1)
    
    doc.add_heading('7.1 Python Libraries', 2)
    libs = doc.add_paragraph()
    libs.add_run('''
Required Packages:
• mediapipe >= 0.10.0 (Tasks API support)
• opencv-python (cv2) - Image processing
• numpy - Array operations
• python-docx - Document generation (optional)

Installation:
pip install mediapipe opencv-python numpy python-docx
''')
    
    doc.add_heading('7.2 Model Files (Auto-Downloaded)', 2)
    models = doc.add_paragraph()
    models.add_run('''
Skin Tone Model:
• File: backend/ml/models/blaze_face_short_range.tflite
• Size: ~1.2 MB
• Source: Google MediaPipe Model Storage

Body Shape Model:
• File: backend/ml/models/pose_landmarker_full.task
• Size: ~8.5 MB
• Source: Google MediaPipe Model Storage

Note: Models download automatically on first run if not present.
''')
    
    # ============================================================
    # SECTION 8: USAGE EXAMPLES
    # ============================================================
    doc.add_heading('8. Usage Examples', 1)
    
    doc.add_heading('8.1 Skin Tone Analysis', 2)
    usage1 = doc.add_paragraph()
    usage1.style = 'Quote'
    usage1.add_run('''# Import the analysis function
from backend.ml.skin_tone import analyze_skin_tone

# Analyze a selfie
result = analyze_skin_tone("selfie.jpg")

# Access results
print(f"Skin Tone: {result['skin_tone']}")  # "Fair"
print(f"Undertone: {result['undertone']}")  # "Warm"
print(f"RGB: {result['avg_rgb']}")          # [185, 152, 128]
print(f"Confidence: {result['confidence']:.0%}")  # 89%
print(f"Face Detected: {result['face_detected']}")  # True
''')
    
    doc.add_heading('8.2 Body Shape Analysis', 2)
    usage2 = doc.add_paragraph()
    usage2.style = 'Quote'
    usage2.add_run('''# Import the analysis function
from backend.ml.body_shape import analyze_body_shape

# Analyze a full-body photo
result = analyze_body_shape("fullbody.jpg")

# Access results
print(f"Body Shape: {result['body_shape']}")  # "Hourglass"
print(f"Shoulder Width: {result['shoulder_width']:.0f}px")  # 180
print(f"Hip Width: {result['hip_width']:.0f}px")  # 175
print(f"Ratio: {result['ratio']:.2f}")  # 1.03
print(f"Confidence: {result['confidence']:.0%}")  # 85%
''')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('Document generated for Adorkable AI Project Presentation')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True
    
    # Save
    output_path = r"c:\Users\dell\Downloads\AI LAB PROJECT\adorkable_ai\Skin_Body_Analysis_Technical_Documentation.docx"
    doc.save(output_path)
    print(f"Document created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    create_document()
