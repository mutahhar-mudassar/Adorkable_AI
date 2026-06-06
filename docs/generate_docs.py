"""
Script to generate all four project documentation .docx files.
Run from the adorkable_ai directory:
    python docs/generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


# =============================================================================
# HELPERS
# =============================================================================

def add_title_page(doc, title, subtitle=""):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph(title)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.runs[0]
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
    if subtitle:
        s = doc.add_paragraph(subtitle)
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.runs[0].font.size = Pt(14)
    doc.add_paragraph()
    meta = doc.add_paragraph(
        "Adorkable AI — AI-Powered Personal Fashion Stylist\n"
        "Academic Year 2025–2026\n"
        "Department of Computer Science & Artificial Intelligence"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.size = Pt(12)
    doc.add_page_break()


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    return p


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F0F0")
    p._p.pPr.append(shading)
    return p


# =============================================================================
# FILE 1: PROJECT ABSTRACT & DETAILS
# =============================================================================

def build_file1():
    doc = Document()
    add_title_page(
        doc,
        "Adorkable AI — Project Abstract & Details",
        "File 1 of 4 | Formal Project Documentation"
    )

    # ---- ABSTRACT ----
    h1(doc, "Abstract")
    body(doc,
        "Adorkable AI is an intelligent, web-based personal fashion styling platform that leverages "
        "computer vision, machine learning, and rule-based artificial intelligence to deliver "
        "context-aware clothing recommendations tailored to each user's unique physical attributes, "
        "wardrobe composition, current weather conditions, and cultural style preferences. The system "
        "integrates MediaPipe's real-time pose estimation and face detection frameworks with a "
        "multi-dimensional outfit scoring engine to analyze a user's skin tone and body shape from "
        "uploaded photographs, then cross-references this biometric data with color theory principles, "
        "weather forecasts obtained via the OpenWeatherMap API, and a pre-curated starter wardrobe "
        "catalog of thirty-two culturally inclusive garments."
    )
    body(doc,
        "At its core, the platform employs a stochastic outfit selector that applies softmax-weighted "
        "random sampling over a scored candidate pool, ensuring recommendations are both high-quality "
        "and varied across sessions. The scoring pipeline evaluates five independent dimensions—color "
        "harmony (30 points), skin tone flattery (20 points), body shape suitability (20 points), "
        "weather appropriateness (20 points), and occasion matching (10 points)—and awards a five-point "
        "trending bonus when garments align with current seasonal fashion trends, yielding a maximum "
        "score of 105 per outfit."
    )
    body(doc,
        "The system supports both Western and Eastern (traditional/ethnic) style preferences, "
        "enforcing strict category pairing rules so that traditional tops are exclusively matched "
        "with traditional bottoms, and includes mandatory hijab recommendations for female users "
        "who follow Eastern styling. A seven-day weekly planner module generates complete outfit "
        "plans with a two-day garment reuse cooldown to maximise wardrobe utilisation, while a "
        "Smart Combo generator produces combinatorial outfit pairings ranked by composite score. "
        "The platform is built with FastAPI for the asynchronous backend and Streamlit for the "
        "interactive frontend, communicating over a RESTful API secured by JWT authentication. "
        "All data is persisted in an SQLite database managed through SQLAlchemy 2.0 with async "
        "support, making the system fully deployable on standard consumer hardware with no cloud "
        "dependency."
    )

    doc.add_page_break()

    # ---- INTRODUCTION ----
    h1(doc, "1. Introduction")
    body(doc,
        "The fashion industry has undergone a significant digital transformation in recent years, "
        "with e-commerce and AI-driven personalisation becoming central to how consumers discover "
        "and engage with clothing. Despite this, the majority of personal wardrobes remain "
        "underutilised—studies indicate that the average person wears only 20% of their clothes "
        "regularly. The challenge is not a lack of clothing, but a lack of intelligent tools to "
        "coordinate existing garments into cohesive, context-appropriate outfits."
    )
    body(doc,
        "Adorkable AI addresses this gap by acting as a personal AI stylist that lives inside "
        "a user's own wardrobe. Unlike e-commerce recommendation systems that push users toward "
        "new purchases, Adorkable AI focuses exclusively on making better use of what users "
        "already own. It analyses clothing images, learns user body attributes from photos, "
        "consults real-time weather data, and applies well-established colour theory and fashion "
        "rules to produce recommendations that are both scientifically grounded and practically "
        "wearable."
    )
    body(doc,
        "The project was developed as a full-stack artificial intelligence application, "
        "demonstrating integration of computer vision, machine learning, rule-based expert "
        "systems, RESTful API design, asynchronous database management, and modern web "
        "frontend development—all within a single cohesive Python ecosystem."
    )

    doc.add_page_break()

    # ---- PROBLEM STATEMENT ----
    h1(doc, "2. Problem Statement")
    body(doc,
        "Current wardrobe management and fashion recommendation tools suffer from several "
        "critical limitations:"
    )
    bullet(doc, "Generic recommendations that ignore the user's actual physical attributes (skin tone, body shape).")
    bullet(doc, "No integration with real-world context such as current weather or temperature.")
    bullet(doc, "Failure to respect cultural and style diversity, particularly for users preferring Eastern or traditional clothing.")
    bullet(doc, "Repetitive outfit suggestions that do not vary across sessions, leading to user disengagement.")
    bullet(doc, "Lack of wardrobe utilisation planning—no weekly coordination or garment reuse tracking.")
    bullet(doc, "No hijab coordination support for users who observe Islamic dress codes.")
    bullet(doc, "Fragmented tools requiring separate apps for wardrobe management, weather checking, and outfit planning.")
    body(doc,
        "Adorkable AI solves each of these problems within a single integrated platform by combining "
        "AI-driven analysis with personalised rule systems and real-time data sources."
    )

    doc.add_page_break()

    # ---- OBJECTIVES ----
    h1(doc, "3. Objectives")
    h2(doc, "3.1 Primary Objectives")
    bullet(doc, "Develop an AI system capable of analysing user skin tone and body shape from uploaded photographs.")
    bullet(doc, "Build a multi-dimensional outfit scoring engine integrating colour theory, body shape rules, weather data, and fashion trends.")
    bullet(doc, "Create a stochastic recommendation engine that balances quality with variety.")
    bullet(doc, "Implement a weekly outfit planner that respects garment reuse constraints.")
    bullet(doc, "Support both Western and Eastern style preferences with appropriate pairing rules.")
    bullet(doc, "Provide mandatory hijab recommendations for female users who select Eastern styling.")

    h2(doc, "3.2 Secondary Objectives")
    bullet(doc, "Deliver a clean, user-friendly web interface accessible without technical expertise.")
    bullet(doc, "Ensure the system functions fully offline except for the optional weather API.")
    bullet(doc, "Design the architecture to be extensible for future features such as purchase recommendations or social sharing.")
    bullet(doc, "Produce comprehensive technical documentation suitable for academic submission and future development.")

    doc.add_page_break()

    # ---- SCOPE ----
    h1(doc, "4. Scope of the Project")
    h2(doc, "4.1 In Scope")
    bullet(doc, "User registration, authentication, and profile management.")
    bullet(doc, "Garment upload with AI-powered colour extraction (dominant colour, hex code, fabric weight classification).")
    bullet(doc, "Skin tone analysis using MediaPipe BlazeFace and cheek-region RGB sampling.")
    bullet(doc, "Body shape classification using MediaPipe Pose (shoulder-to-hip ratio analysis).")
    bullet(doc, "Five-dimensional outfit scoring: colour harmony, skin flattery, body shape, weather, occasion.")
    bullet(doc, "Daily outfit recommendation with 're-imagine' capability (cycle through top-N alternatives).")
    bullet(doc, "Seven-day weekly planner with two-day garment cooldown.")
    bullet(doc, "Quick Plan (one-click 7-day plan with default occasion pattern).")
    bullet(doc, "Smart Combo generator for combinatorial outfit pairing.")
    bullet(doc, "Hijab rotation and mandatory selection for Eastern-styled female users.")
    bullet(doc, "Analytics dashboard showing wardrobe utilisation, wear frequency, and colour distribution.")
    bullet(doc, "Starter wardrobe catalog of 32 gender-filtered preloaded garments.")

    h2(doc, "4.2 Out of Scope")
    bullet(doc, "E-commerce integration or product purchase links.")
    bullet(doc, "Social features (sharing, following, community feeds).")
    bullet(doc, "Mobile native application (iOS/Android).")
    bullet(doc, "Integration with physical smart wardrobe hardware.")
    bullet(doc, "Real-time video try-on or augmented reality features.")

    doc.add_page_break()

    # ---- SIGNIFICANCE ----
    h1(doc, "5. Significance of the Project")
    h2(doc, "5.1 Academic Significance")
    body(doc,
        "This project demonstrates the practical application of multiple AI and software engineering "
        "disciplines within a single deployable system. It provides hands-on implementation of "
        "computer vision pipelines (MediaPipe), colour science (HSL colour space, harmony rules), "
        "probabilistic selection algorithms (softmax-weighted sampling), RESTful API design "
        "(FastAPI with async/await), relational database management (SQLAlchemy ORM), and "
        "modern web frontend development (Streamlit)."
    )
    h2(doc, "5.2 Social Significance")
    body(doc,
        "Adorkable AI is designed with cultural inclusivity at its core. By supporting traditional "
        "Eastern garments (salwar kameez, shalwar, dupatta) alongside Western clothing, and by "
        "including automatic hijab coordination for observant Muslim users, the platform serves a "
        "broader and more diverse user base than most existing Western-centric fashion applications. "
        "This inclusivity is not superficial—it is enforced at the algorithmic level through "
        "dedicated category types (traditional_top, traditional_bottom) and pairing rules."
    )
    h2(doc, "5.3 Environmental Significance")
    body(doc,
        "By optimising use of existing wardrobes rather than encouraging new purchases, "
        "Adorkable AI supports sustainable fashion practices. The weekly planner's garment "
        "rotation logic ensures users rediscover underutilised clothing, directly countering "
        "the fast-fashion consumption pattern."
    )

    doc.add_page_break()

    # ---- METHODOLOGY ----
    h1(doc, "6. Methodology Overview")
    h2(doc, "6.1 System Architecture")
    body(doc,
        "Adorkable AI follows a client-server architecture with three main layers:"
    )
    bullet(doc, "Frontend Layer: Streamlit web application (pages 1–7) communicating with the backend via HTTP/REST.")
    bullet(doc, "Backend Layer: FastAPI application exposing RESTful endpoints at /api/v1/, handling authentication, data management, and AI inference.")
    bullet(doc, "Data Layer: SQLite database accessed asynchronously via SQLAlchemy 2.0 and aiosqlite.")

    h2(doc, "6.2 AI Pipeline")
    body(doc, "The AI pipeline processes user data in four stages:")
    bullet(doc, "Stage 1 – Profile Analysis: User uploads a selfie (skin tone) and full-body photo (body shape). MediaPipe models extract facial and skeletal landmarks. Skin tone is classified across 6 levels with 3 undertone categories. Body shape is classified into 6 types based on shoulder-hip-waist ratios.")
    bullet(doc, "Stage 2 – Wardrobe Analysis: Each uploaded garment image is processed by OpenCV's KMeans colour extraction to identify the dominant colour and its hex code. Garment metadata (category, style, fabric weight, occasion tags) is stored in the database.")
    bullet(doc, "Stage 3 – Candidate Generation: For a given occasion, weather, and style preference, all valid top+bottom and dress combinations are generated from the wardrobe. Eastern style enforces traditional_top × traditional_bottom pairing only.")
    bullet(doc, "Stage 4 – Scoring & Selection: Each candidate outfit is scored across five dimensions. The top-N candidates are passed through a softmax-weighted random selector to choose the final recommendation, ensuring quality with variety.")

    h2(doc, "6.3 Development Methodology")
    body(doc,
        "The project was developed using an iterative Agile approach with short development "
        "sprints focused on: authentication → wardrobe → AI profile analysis → outfit scoring "
        "→ planner → smart combo → frontend polish → documentation. Each feature was "
        "independently testable via the FastAPI /docs Swagger interface before frontend "
        "integration."
    )

    doc.save(os.path.join(OUT_DIR, "1_Project_Abstract_Details.docx"))
    print("✅ File 1 saved.")


# =============================================================================
# FILE 2: CODE LINE-BY-LINE EXPLANATION
# =============================================================================

def build_file2():
    doc = Document()
    add_title_page(
        doc,
        "Adorkable AI — Code Line-by-Line Explanation",
        "File 2 of 4 | Complete Source Code Documentation"
    )

    # ===== BACKEND/MAIN.PY =====
    h1(doc, "1. backend/main.py — Application Entry Point")
    body(doc,
        "This is the root file that creates and configures the FastAPI application, "
        "registers all routers, and manages the application lifecycle."
    )
    h2(doc, "Imports and Configuration")
    code_block(doc, "from contextlib import asynccontextmanager")
    body(doc, "asynccontextmanager converts a generator function into an async context manager, used to define startup/shutdown hooks for the FastAPI app.")
    code_block(doc, "from fastapi.middleware.cors import CORSMiddleware")
    body(doc, "Enables Cross-Origin Resource Sharing so the Streamlit frontend (port 8501) can call the FastAPI backend (port 8006) without browser security blocks.")
    code_block(doc, "from fastapi.staticfiles import StaticFiles")
    body(doc, "Mounts the /uploads directory as a static file server, allowing garment images stored on disk to be served directly via URL.")

    h2(doc, "Lifespan Handler")
    code_block(doc, "@asynccontextmanager\nasync def lifespan(app: FastAPI):\n    await create_all_tables()\n    os.makedirs(UPLOAD_DIR, exist_ok=True)\n    yield  # Application runs here\n    # Shutdown code after yield")
    body(doc, "The lifespan function runs code before (startup) and after (shutdown) the application. create_all_tables() ensures all SQLAlchemy ORM tables exist in SQLite on first run.")

    h2(doc, "Router Registration")
    code_block(doc, "app.include_router(auth_router, prefix='/api/v1')\napp.include_router(wardrobe.router, prefix='/api/v1')\napp.include_router(recommendations.router, prefix='/api/v1')")
    body(doc, "Each feature module (auth, wardrobe, profile, recommendations, planner, combo, analytics) is a separate router registered with the /api/v1 prefix, resulting in URLs like /api/v1/wardrobe/, /api/v1/recommend, etc.")

    doc.add_page_break()

    # ===== BACKEND/DATABASE.PY =====
    h1(doc, "2. backend/database.py — Data Models & Database Layer")
    body(doc, "Defines all SQLAlchemy ORM models and async database utility functions.")

    h2(doc, "Async Engine Setup")
    code_block(doc, "ASYNC_DATABASE_URL = DATABASE_URL.replace('sqlite:///', 'sqlite+aiosqlite:///')\nengine = create_async_engine(ASYNC_DATABASE_URL, echo=False, future=True)")
    body(doc, "SQLAlchemy's standard sqlite:/// driver is synchronous. By replacing it with sqlite+aiosqlite:///, we enable non-blocking async database operations essential for a high-performance FastAPI application.")

    h2(doc, "User Model")
    code_block(doc, "class User(Base):\n    id: Mapped[int] = mapped_column(Integer, primary_key=True)\n    email: Mapped[str] = mapped_column(String(255), unique=True)\n    hashed_password: Mapped[str] = mapped_column(String(255))\n    gender: Mapped[Optional[str]]\n    city: Mapped[Optional[str]]")
    body(doc, "The User model stores authentication credentials and basic profile info. gender is used to filter starter wardrobe items and to enable hijab recommendations. city is used to fetch location-specific weather forecasts.")

    h2(doc, "GarmentItem Model")
    code_block(doc, "class GarmentItem(Base):\n    category: str       # top, bottom, dress, outerwear, traditional_top,\n                        # traditional_bottom, hijab\n    dominant_color: str # e.g. 'Navy Blue'\n    color_hex: str      # e.g. '#000080'\n    fabric_weight: str  # light, medium, heavy\n    occasion_tags: str  # JSON string e.g. '[\"Casual\", \"Work\"]'\n    wear_count: int     # increments when outfit is worn")
    body(doc, "The category field is central to outfit generation logic. The engine uses it to determine valid outfit combinations. occasion_tags is stored as a JSON string due to SQLite's lack of native array support; the get_occasion_tags_list() method deserialises it.")

    h2(doc, "OutfitLog Model")
    code_block(doc, "class OutfitLog(Base):\n    top_id: Optional[int]     # nullable FK - could be a dress outfit\n    bottom_id: Optional[int]  # nullable FK\n    dress_id: Optional[int]   # nullable FK\n    score: float\n    trending_badge: bool\n    worn_date: date")
    body(doc, "Logs every outfit recommendation that is generated. All garment FKs are nullable because an outfit may be a dress-only combination (no top/bottom). Used by the analytics module to calculate wear frequency and colour distribution.")

    h2(doc, "get_db() Dependency")
    code_block(doc, "async def get_db():\n    async with async_session_maker() as session:\n        try:\n            yield session\n            await session.commit()\n        except Exception:\n            await session.rollback()\n            raise")
    body(doc, "This is a FastAPI dependency generator. When used with Depends(get_db), FastAPI automatically opens a new database session per request, commits it on success, rolls it back on exception, and closes it after the response is sent. This prevents connection leaks.")

    doc.add_page_break()

    # ===== BACKEND/AUTH.PY =====
    h1(doc, "3. backend/auth.py — Authentication System")

    h2(doc, "Password Hashing")
    code_block(doc, "pwd_context = CryptContext(schemes=['bcrypt'], deprecated='auto')\n\ndef hash_password(plain): return pwd_context.hash(plain)\ndef verify_password(plain, hashed): return pwd_context.verify(plain, hashed)")
    body(doc, "Passwords are never stored in plain text. bcrypt hashing applies a random salt and multiple rounds of hashing, making rainbow table attacks infeasible. verify_password re-hashes the input and compares bit-for-bit.")

    h2(doc, "JWT Token Creation")
    code_block(doc, "def create_access_token(data: dict) -> str:\n    to_encode = data.copy()\n    to_encode['sub'] = str(to_encode['sub'])  # RFC-compliant string\n    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)\n    to_encode.update({'exp': expire, 'iat': datetime.utcnow(), 'type': 'access'})\n    return jwt.encode(to_encode, SECRET_KEY, algorithm='HS256')")
    body(doc, "JWT (JSON Web Token) encodes the user's ID as the 'sub' (subject) claim. The token is signed with the SECRET_KEY using HS256 algorithm. Tokens expire after 7 days (configured in config.py). 'iat' (issued at) is included for audit purposes.")

    h2(doc, "get_current_user Dependency")
    code_block(doc, "async def get_current_user(credentials: HTTPAuthorizationCredentials, db):\n    token = credentials.credentials  # Extract from 'Bearer <token>'\n    payload = decode_token(token)    # Verify signature & expiry\n    user_id = payload.get('sub')     # Get user ID\n    user = await get_user_by_id(db, int(user_id))  # DB lookup\n    return user")
    body(doc, "Every protected API endpoint uses Depends(get_current_user) to authenticate requests. The Bearer token is extracted from the Authorization header, decoded and verified, and the corresponding User object is fetched from the database and injected into the route handler.")

    doc.add_page_break()

    # ===== ENGINE/OUTFIT_SCORER.PY =====
    h1(doc, "4. backend/engine/outfit_scorer.py — Core Scoring Engine")
    body(doc, "The most critical file in the project. Calculates a composite score for any garment combination.")

    h2(doc, "Score Dimensions")
    code_block(doc, "# SCORING_WEIGHTS from config.py:\n# color_harmony: 30\n# skin_flattery: 20\n# body_shape:    20\n# weather:       20\n# occasion:      10\n# trending bonus: 5\n# TOTAL MAX:    105")
    body(doc, "The weights reflect the relative importance of each factor. Colour harmony is given the highest weight because it is the most visually obvious element of a well-coordinated outfit.")

    h2(doc, "score_outfit() Master Function")
    code_block(doc, "def score_outfit(garments, user_profile, weather, occasion, style_pref):\n    color_score, _ = score_color_harmony(garments)\n    skin_score, _  = score_skin_tone_flattery(garments, skin_tone, undertone)\n    body_score, _  = score_body_shape_suitability(garments, body_shape)\n    weather_score  = avg(weather_suitability_score(g, temp, cond) for g in garments)\n    occasion_score, _ = score_occasion_match(garments, occasion)\n    base_score = weighted_sum_of_all\n    if is_trending: base_score += 5\n    return {score, top, bottom, dress, outerwear, why_this_suits_you, ...}")
    body(doc, "Each sub-scorer returns a normalised float (0.0–1.0), which is multiplied by its weight to yield a contribution in points. The final score is returned as part of a rich dictionary that includes all garment references and textual explanations for UI display.")

    h2(doc, "check_trending()")
    code_block(doc, "def check_trending(outfit_colors, outfit_style, categories):\n    season = get_current_season()   # spring_summer / autumn_winter\n    trends = load_trends(season)    # from backend/data/trends.json\n    # Check color matches, style matches, eastern/western trend matches\n    return (is_trending: bool, trending_reason: str)")
    body(doc, "Trend data is stored in a JSON file (trends.json) with seasonal keys like '2026_spring_summer'. The function cross-references outfit colours and style against trending colours and styles for the current season. A match awards the 5-point trending bonus.")

    doc.add_page_break()

    # ===== ENGINE/COLOR_THEORY.PY =====
    h1(doc, "5. backend/engine/color_theory.py — Colour Science Engine")

    h2(doc, "Colour Space Conversion")
    code_block(doc, "def hex_to_hsl(hex_str: str) -> Tuple[float, float, float]:\n    r, g, b = hex_to_rgb(hex_str)\n    h, l, s = colorsys.rgb_to_hls(r/255, g/255, b/255)\n    return (h*360, s*100, l*100)  # degrees, percent, percent")
    body(doc, "Colour harmony calculations work in HSL (Hue-Saturation-Lightness) space rather than RGB because HSL intuitively maps to how humans perceive colour relationships. Two colours 180° apart on the hue wheel are complementary; colours within 30° are analogous.")

    h2(doc, "Harmony Rules")
    code_block(doc, "def is_complementary(h1, h2):   return abs(h1 - h2) in range(150, 210)\ndef is_analogous(h1, h2):       return abs(h1 - h2) <= 30\ndef is_monochromatic(hex1,hex2): return same hue ± 15°, different lightness\ndef is_triadic(h1, h2):         return abs(h1 - h2) close to 120°")
    body(doc, "These are the four classical colour harmony rules from Itten's colour wheel theory, as used in art, design, and fashion education. The scoring engine awards higher points for complementary and triadic combinations (high contrast, high visual impact) and moderate points for analogous and monochromatic (subtle, elegant).")

    h2(doc, "Skin Flattery")
    code_block(doc, "def is_flattering_for_skin(color_hex, skin_tone, undertone):\n    # Warm undertone: flatters warm colours (coral, gold, olive, rust)\n    # Cool undertone: flatters cool colours (navy, purple, rose, emerald)\n    # Neutral: flatters most colours")
    body(doc, "This function implements the concept of 'personal colour seasons' from professional colour analysis. The system maps skin undertones to flattering colour families and awards points when an outfit's dominant colours align with the user's undertone.")

    doc.add_page_break()

    # ===== ENGINE/OUTFIT_CONSTRAINTS.PY =====
    h1(doc, "6. backend/engine/outfit_constraints.py — Outfit Rules Engine")

    h2(doc, "Eastern Style Pairing Enforcement")
    code_block(doc, "for top in tops:\n    for bottom in bottoms:\n        if style_pref == 'Eastern':\n            if not ('traditional' in top.category\n                    and 'traditional' in bottom.category):\n                continue  # Skip — Western pieces excluded")
    body(doc, "This is the enforced Eastern style rule. When a user selects Eastern style, the engine only considers combinations where BOTH the top and bottom are from traditional categories. This prevents mixing Western jeans with a kurta, for example, when the user explicitly wants Eastern outfits.")

    h2(doc, "Completeness Rule")
    code_block(doc, "def outfit_has_complete_separates_or_dress(scored):\n    if scored.get('dress'):               # Dress = complete outfit alone\n        return not(top or bottom)         # No mixing dress + separates\n    return scored['top'] and scored['bottom']  # Must have both")
    body(doc, "A dress is treated as a complete outfit by itself. The function enforces that no outfit can contain a dress AND separates simultaneously (which would be nonsensical). A top-only or bottom-only outfit is also invalid.")

    h2(doc, "Weather Outerwear Rule")
    code_block(doc, "def outfit_passes_weather_outerwear_rule(scored, weather, ow_avail):\n    if not ow_avail: return True          # No outerwear in wardrobe, skip rule\n    if not needs_outerwear(weather): return True   # Warm weather, skip\n    return scored.get('outerwear') is not None  # Cold+owned → must include")
    body(doc, "If the temperature is below threshold AND the user owns outerwear, the engine only accepts outfits that include that outerwear. This prevents cold-weather recommendations that ignore available jackets and coats.")

    doc.add_page_break()

    # ===== ENGINE/STOCHASTIC_SELECTOR.PY =====
    h1(doc, "7. backend/engine/stochastic_selector.py — Probabilistic Selection")

    h2(doc, "softmax() Function")
    code_block(doc, "def softmax(weights):\n    max_w = max(weights)                   # Subtract max for stability\n    exp_w = [math.exp(w - max_w) for w in weights]\n    return [e / sum(exp_w) for e in exp_w] # Probabilities sum to 1")
    body(doc, "Softmax converts raw scores into probabilities. Subtracting the maximum value before exponentiation prevents floating-point overflow. The result is a probability distribution where higher-scored outfits have higher selection probability but lower-scored ones still have a non-zero chance.")

    h2(doc, "weighted_random_select()")
    code_block(doc, "def weighted_random_select(scored_outfits, top_n=5):\n    sorted_outfits = sorted(outfits, key=score, reverse=True)\n    candidates = sorted_outfits[:top_n]    # Consider only top 5\n    weights = softmax([o['score'] for o in candidates])\n    return random.choices(candidates, weights=weights, k=1)[0]")
    body(doc, "The function only considers the top-N candidates (default 5) to maintain quality, but uses softmax weighting for selection so the algorithm does not always return the #1 scored outfit. This creates variety across sessions while guaranteeing that only good outfits are shown.")

    h2(doc, "select_with_exclusion()")
    code_block(doc, "def select_with_exclusion(outfits, used_item_ids):\n    available = [o for o in outfits\n                 if not garment_ids(o).intersection(used_item_ids)]\n    return weighted_random_select(available) if available else None")
    body(doc, "Used by the weekly planner to enforce the 2-day garment cooldown. Outfits containing any recently-used garment ID are filtered out before selection.")

    doc.add_page_break()

    # ===== ML/SKIN_TONE.PY =====
    h1(doc, "8. backend/ml/skin_tone.py — Skin Tone Analysis")

    h2(doc, "MediaPipe Face Detection")
    code_block(doc, "FACE_DETECTOR = FaceDetector.create_from_options(options)\n# Model: blaze_face_short_range.tflite (auto-downloaded)")
    body(doc, "BlazeFace is a lightweight face detection model designed by Google for mobile/edge devices. It detects facial bounding boxes with sub-millisecond latency. The .tflite (TensorFlow Lite) model is downloaded automatically on first run if not present.")

    h2(doc, "Cheek Region Sampling")
    code_block(doc, "# After detecting face bounding box:\n# Sample pixels from left/right cheek regions (avoid nose, forehead)\n# Average the RGB values of these cheek pixels\n# Map average_brightness to one of 6 skin tone categories\n# Map R/G ratio to undertone (Warm/Cool/Neutral)")
    body(doc, "The cheek region is chosen because it provides the most consistent skin colour sample—the nose and forehead often have highlights or shadows. The R/G ratio is a proxy for warm vs. cool undertones: higher red component indicates warmth.")

    h2(doc, "6-Level Classification")
    code_block(doc, "SKIN_TONE_LABELS = [\n    'Very Fair',    # Porcelain/Ivory\n    'Fair',         # Light\n    'Light Medium', # Olive/Light Tan\n    'Medium',       # Tan/Golden\n    'Medium Dark',  # Brown/Caramel\n    'Dark'          # Deep/Rich\n]")
    body(doc, "The 6-level scale (based on the Fitzpatrick scale concept) provides finer granularity than simple Fair/Medium/Dark, enabling more precise colour palette recommendations. Each level maps to different flattering colour palettes in the skin tone palette JSON file.")

    doc.add_page_break()

    # ===== ML/BODY_SHAPE.PY =====
    h1(doc, "9. backend/ml/body_shape.py — Body Shape Analysis")

    h2(doc, "MediaPipe Pose Landmarker")
    code_block(doc, "POSE_LANDMARKER = PoseLandmarker.create_from_options(options)\n# Model: pose_landmarker_full.task\n# Detects 33 body landmarks in 3D space (x, y, z)")
    body(doc, "MediaPipe's PoseLandmarker model detects 33 anatomical landmarks on the human body including shoulders, hips, elbows, wrists, and knees. The 'full' model variant is used for higher accuracy (vs. 'lite') at moderate computational cost.")

    h2(doc, "Measurement Extraction")
    code_block(doc, "# Key landmark indices:\nLEFT_SHOULDER  = 11;  RIGHT_SHOULDER = 12\nLEFT_HIP       = 23;  RIGHT_HIP      = 24\n\nshoulder_width = pixel_distance(left_shoulder, right_shoulder)\nhip_width      = pixel_distance(left_hip, right_hip)\nratio = shoulder_width / hip_width")
    body(doc, "Body shape is determined mathematically from the ratio of shoulder width to hip width. Additional metrics like waist estimation (midpoint between shoulders and hips) refine the classification. All measurements are taken in normalised image coordinates to be scale-invariant.")

    h2(doc, "Shape Classification Rules")
    code_block(doc, "if ratio > 1.15:      shape = 'Inverted Triangle'\nelif ratio < 0.85:    shape = 'Pear'\nelif 0.95 < ratio < 1.05 and waist_narrow: shape = 'Hourglass'\nelif waist_similar:   shape = 'Rectangle'\nelif waist_full:      shape = 'Apple'\nelse:                 shape = 'Athletic'")
    body(doc, "The thresholds are based on standard fashion industry body type classification criteria. Each shape maps to specific clothing recommendations in BODY_SHAPE_RULES (defined in config.py) that favour certain silhouettes and discourage others.")

    doc.add_page_break()

    # ===== ENGINE/WEEKLY_PLANNER.PY =====
    h1(doc, "10. backend/engine/weekly_planner.py — Weekly Planner Engine")

    h2(doc, "7-Day Loop with Cooldown")
    code_block(doc, "used_items_with_day = []  # [(item_id, day_index), ...]\nCOOLDOWN_DAYS = 2\n\nfor day_idx in range(7):\n    recently_used = [\n        item_id for item_id, used_day in used_items_with_day\n        if day_idx - used_day < COOLDOWN_DAYS\n    ]\n    selected = select_with_exclusion(candidates, recently_used)")
    body(doc, "The cooldown prevents garments used on day 0 from appearing again on day 1, but allows them on day 2. This is a standard wardrobe rotation strategy: items worn on Monday can be re-worn on Wednesday. The list of (item_id, day_index) tuples allows fine-grained tracking.")

    h2(doc, "Bottom Variety Boosting")
    code_block(doc, "def _boost_unused_bottoms(candidates, used_bottom_ids):\n    for outfit in candidates:\n        if outfit['bottom'].id not in used_bottom_ids:\n            outfit['score'] += 15.0  # Boost for unused bottom")
    body(doc, "Bottoms are the slowest-rotating item in most wardrobes. The +15 score boost for outfits featuring unused bottoms strongly encourages the engine to pick different bottoms each day, without hard-excluding already-used ones.")

    h2(doc, "Actual Weather per Day")
    code_block(doc, "forecast = get_7day_forecast(city)  # 7-element list from weather API\nfor day_idx in range(7):\n    weather = forecast[day_idx]\n    # Override weather to 22°C only for candidate generation (flexible)\n    # Store actual weather in selected outfit for display")
    body(doc, "The engine uses a neutral 22°C for candidate scoring (to prevent weather from eliminating too many garments) but stores the actual forecast temperature in the day's plan for accurate weather explanation display in the frontend.")

    doc.add_page_break()

    # ===== FRONTEND =====
    h1(doc, "11. frontend/ — Streamlit Frontend Pages")

    h2(doc, "Page Structure")
    bullet(doc, "app.py — Landing page and navigation entry point.")
    bullet(doc, "1_Register_Login.py — User authentication (register/login/logout).")
    bullet(doc, "2_My_Profile.py — Profile management, skin tone & body shape upload.")
    bullet(doc, "3_My_Wardrobe.py — Garment upload, categorisation, wardrobe grid view.")
    bullet(doc, "4_Daily_Outfit.py — Daily outfit recommendation with re-imagine.")
    bullet(doc, "5_Weekly_Planner.py — 7-day plan display with garment images and weather.")
    bullet(doc, "6_Smart_Combo.py — Combinatorial outfit generator.")
    bullet(doc, "7_Analytics.py — Wardrobe utilisation charts and statistics.")

    h2(doc, "API Communication Pattern")
    code_block(doc, "headers = {'Authorization': f'Bearer {st.session_state.user_token}'}\nresponse = httpx.post(\n    f'{API_BASE}/recommend',\n    headers=headers,\n    json={'occasion': occasion, 'style_pref': style_pref},\n    timeout=httpx.Timeout(120.0)\n)")
    body(doc, "The frontend uses httpx (async-capable HTTP client) to communicate with the backend. JWT tokens are stored in st.session_state and sent as Bearer tokens on every authenticated request. A 120-second timeout accommodates the compute time of AI inference.")

    h2(doc, "resolve_image_path Helper")
    code_block(doc, "def resolve_image_path(path: str) -> Optional[str]:\n    if not path: return None\n    if os.path.exists(path): return path\n    # Try relative to project root\n    abs_path = os.path.join(PROJECT_ROOT, path.lstrip('/'))\n    return abs_path if os.path.exists(abs_path) else None")
    body(doc, "Image paths stored in the database may be absolute or relative depending on how garments were uploaded. This function tries multiple path resolution strategies to always locate the correct image file for display.")

    doc.save(os.path.join(OUT_DIR, "2_Code_Line_by_Line_Explanation.docx"))
    print("✅ File 2 saved.")


# =============================================================================
# FILE 3: LIBRARIES, COMPONENTS, ARCHITECTURE
# =============================================================================

def build_file3():
    doc = Document()
    add_title_page(
        doc,
        "Adorkable AI — Libraries, Components & System Architecture",
        "File 3 of 4 | Technical Stack Deep Dive"
    )

    h1(doc, "1. Complete Library Reference")

    libs = [
        ("FastAPI 0.115", "Backend web framework",
         "FastAPI is a modern Python web framework built on Starlette and Pydantic. It uses Python type hints to automatically generate OpenAPI documentation and validate request/response data. Its async-first design (built on ASGI/uvicorn) makes it ideal for handling concurrent AI inference requests without blocking. Chosen over Flask/Django for its performance, automatic docs, and native async support."),
        ("Uvicorn 0.32", "ASGI Server",
         "Uvicorn is a lightning-fast ASGI (Asynchronous Server Gateway Interface) server that runs the FastAPI application. It handles HTTP/1.1 connections, manages worker processes, and supports hot-reload during development. The [standard] variant includes performance extras like uvloop and httptools."),
        ("Streamlit 1.40", "Frontend Framework",
         "Streamlit converts pure Python scripts into interactive web applications without requiring HTML/CSS/JavaScript knowledge. Each page is a .py file that renders UI elements (st.image, st.button, st.columns) declaratively. Chosen for rapid prototyping and its built-in reactive state management."),
        ("SQLAlchemy 2.0 + aiosqlite", "Database ORM & Async Driver",
         "SQLAlchemy 2.0 provides the ORM layer using Python dataclasses-style mapped_column syntax. The async_sessionmaker and create_async_engine enable non-blocking database operations. aiosqlite provides the asyncio-compatible SQLite driver. Together, they allow hundreds of concurrent database operations without thread blocking."),
        ("MediaPipe 0.10.18", "Computer Vision AI",
         "MediaPipe is Google's open-source ML framework for real-time media processing. The project uses two models: BlazeFace (face detection for skin tone analysis) and PoseLandmarker (33-point body skeleton for body shape analysis). Both models are TensorFlow Lite format, running on CPU with sub-second inference time."),
        ("OpenCV (cv2) 4.10", "Image Processing",
         "OpenCV (Open Computer Vision Library) provides low-level image processing primitives: reading images, converting colour spaces (BGR to RGB), resizing, cropping, and computing KMeans colour clustering for dominant colour extraction. Headless variant is used (no GUI display needed on server)."),
        ("TensorFlow 2.18 + Keras 3.6", "Deep Learning Backend",
         "TensorFlow serves as the computational backend for MediaPipe's TFLite models. Keras provides the high-level neural network API. While not directly used for custom model training in the final version, the infrastructure supports future clothing classifier training."),
        ("scikit-learn 1.5.2", "Machine Learning Utilities",
         "Used for KMeans clustering in the colour extraction pipeline. When a garment image is uploaded, the KMeans algorithm clusters pixel colours into k=5 clusters and identifies the dominant cluster as the garment's main colour."),
        ("NumPy 2.0.2", "Numerical Computing",
         "NumPy's ndarray is the universal image representation format used throughout the pipeline: OpenCV returns BGR arrays, MediaPipe accepts numpy arrays, and colour calculations operate on numpy vectors. Also used for statistical operations in analytics."),
        ("Pillow 11.0", "Image Loading & Saving",
         "Pillow (PIL) handles image file I/O across formats (JPEG, PNG, WebP). Used to open uploaded images before passing to OpenCV/MediaPipe, and to resize/thumbnail images for efficient storage."),
        ("python-jose[cryptography]", "JWT Authentication",
         "Implements JWT (JSON Web Token) generation and verification using the HS256 algorithm. Tokens contain the user's ID as the 'sub' claim and expire after 7 days. The cryptography extra provides production-grade PKCS operations."),
        ("passlib[bcrypt]", "Password Security",
         "Provides bcrypt password hashing with automatic salting. bcrypt is the industry standard for password storage due to its adaptive cost factor, which can be increased as hardware improves to maintain security."),
        ("httpx 0.27", "HTTP Client",
         "The frontend uses httpx to make synchronous HTTP requests to the backend API. httpx was chosen over the standard requests library because it supports both sync and async modes, has better timeout control, and handles HTTP/2."),
        ("python-dotenv 1.0.1", "Environment Configuration",
         "Loads key-value pairs from the .env file into the process environment. This separates configuration (API keys, database URLs) from code, following the 12-factor app methodology."),
        ("plotly 5.24", "Interactive Charts",
         "Renders the analytics dashboard charts (wear frequency bar charts, colour distribution pie charts, category breakdown charts) as interactive HTML/JavaScript visualisations embedded in the Streamlit frontend."),
        ("pydantic (via FastAPI)", "Data Validation",
         "Pydantic BaseModel classes define the schema of every API request and response. FastAPI automatically validates incoming JSON against these schemas, returns 422 Unprocessable Entity on validation failure, and generates OpenAPI documentation from them."),
    ]

    for name, role, explanation in libs:
        h2(doc, f"{name} — {role}")
        body(doc, explanation)

    doc.add_page_break()

    # ---- SYSTEM ARCHITECTURE ----
    h1(doc, "2. System Architecture")

    h2(doc, "2.1 High-Level Architecture")
    body(doc, "Adorkable AI follows a 3-tier architecture:")
    bullet(doc, "Tier 1 (Presentation): Streamlit frontend — 8 pages communicating via httpx REST calls.")
    bullet(doc, "Tier 2 (Application): FastAPI backend — 7 API modules (auth, wardrobe, profile, recommendations, planner, combo, analytics) with AI engine modules.")
    bullet(doc, "Tier 3 (Data): SQLite database (4 tables: users, user_profiles, garment_items, outfit_logs) + flat-file data (preloaded_wardrobe.json, trends.json, skin_tone_palettes.json).")

    h2(doc, "2.2 API Endpoint Map")
    endpoints = [
        ("/api/v1/auth/register", "POST", "User registration with email/password/gender/city"),
        ("/api/v1/auth/login", "POST", "Login, returns JWT token"),
        ("/api/v1/auth/me", "GET", "Get current user info"),
        ("/api/v1/wardrobe/", "GET", "List all garments with optional category filter"),
        ("/api/v1/wardrobe/upload", "POST", "Upload garment image with AI analysis"),
        ("/api/v1/wardrobe/stats", "GET", "Wardrobe statistics (count, by category, by colour)"),
        ("/api/v1/wardrobe/{id}/worn", "POST", "Mark garment as worn (increment wear_count)"),
        ("/api/v1/profile/", "GET", "Get user profile with skin tone & body shape"),
        ("/api/v1/profile/selfie", "POST", "Upload selfie → MediaPipe skin tone analysis"),
        ("/api/v1/profile/body", "POST", "Upload body photo → MediaPipe body shape analysis"),
        ("/api/v1/recommend", "POST", "Generate daily outfit recommendation"),
        ("/api/v1/plan/weekly", "POST", "Generate 7-day weekly plan"),
        ("/api/v1/plan/quick", "GET", "Quick 7-day plan with default occasions"),
        ("/api/v1/combo", "POST", "Generate smart outfit combinations"),
        ("/api/v1/analytics/dashboard", "GET", "Analytics dashboard data"),
    ]
    for ep, method, desc in endpoints:
        bullet(doc, f"[{method}] {ep} — {desc}")

    doc.add_page_break()

    h1(doc, "3. Database Design")
    h2(doc, "Entity Relationship Summary")
    body(doc, "Four tables with the following relationships:")
    bullet(doc, "users (1) ──< garment_items (many): Each user owns many garments. CASCADE DELETE removes all garments when user is deleted.")
    bullet(doc, "users (1) ──< outfit_logs (many): Each user has many outfit logs. CASCADE DELETE applies.")
    bullet(doc, "users (1) ──< user_profiles (1): One-to-one relationship. A user has at most one profile.")
    bullet(doc, "outfit_logs >── garment_items: Each log references up to 4 garments (top, bottom, dress, outerwear) via nullable FKs with SET NULL on delete.")

    h2(doc, "GarmentItem Key Fields")
    code_block(doc,
        "category:       top | bottom | dress | outerwear |\n"
        "                traditional_top | traditional_bottom | hijab\n"
        "style:          casual | formal | business | traditional_eastern | ...\n"
        "fabric_weight:  light | medium | heavy\n"
        "occasion_tags:  JSON list — ['Casual','Work','Party',...]\n"
        "gender_fit:     male | female | unisex"
    )

    doc.add_page_break()

    h1(doc, "4. AI/ML Models Used")

    h2(doc, "4.1 BlazeFace (Face Detection)")
    body(doc, "Model: blaze_face_short_range.tflite")
    bullet(doc, "Type: Convolutional Neural Network (anchor-based single-shot detector)")
    bullet(doc, "Input: 128×128 RGB image")
    bullet(doc, "Output: Bounding boxes with confidence scores for detected faces")
    bullet(doc, "Usage: Detect face bounding box, crop cheek region, sample skin colour")
    bullet(doc, "Source: Google MediaPipe Model Registry (auto-downloaded)")

    h2(doc, "4.2 PoseLandmarker Full (Body Pose)")
    body(doc, "Model: pose_landmarker_full.task")
    bullet(doc, "Type: Multi-stage CNN with heatmap prediction")
    bullet(doc, "Input: RGB image (variable size, internally resized to 256×256)")
    bullet(doc, "Output: 33 landmarks × (x, y, z, visibility) in normalised image coordinates")
    bullet(doc, "Usage: Extract shoulder (11,12) and hip (23,24) landmark positions, compute width ratio")
    bullet(doc, "Source: Google MediaPipe Model Registry (auto-downloaded)")

    h2(doc, "4.3 KMeans Colour Clustering (scikit-learn)")
    bullet(doc, "Type: Unsupervised clustering algorithm")
    bullet(doc, "Input: Flattened pixel array of garment image (n_samples × 3 for RGB)")
    bullet(doc, "Output: k=5 cluster centres, dominant cluster = garment colour")
    bullet(doc, "Usage: Determine dominant_color hex and map to fashion colour name")

    h2(doc, "4.4 Rule-Based Outfit Scorer (Custom)")
    body(doc, "No pre-trained model — the scorer is a deterministic expert system:")
    bullet(doc, "Colour harmony: HSL hue-wheel mathematics (complementary, analogous, etc.)")
    bullet(doc, "Skin flattery: Lookup table mapping skin undertone to flattering colour families")
    bullet(doc, "Body shape: Rule lookup from BODY_SHAPE_RULES config dictionary")
    bullet(doc, "Weather: Temperature threshold rules with humidity modifier")
    bullet(doc, "Occasion: Keyword matching between garment tags and occasion label")

    doc.add_page_break()

    h1(doc, "5. Key Algorithms")

    h2(doc, "5.1 Softmax Weighted Random Selection")
    code_block(doc,
        "# Given scores [85, 82, 78]:\n"
        "exp = [e^(85-85), e^(82-85), e^(78-85)] = [1, 0.05, 0.0009]\n"
        "probs = [0.952, 0.047, 0.0009]\n"
        "# Top-5 selected with these probabilities"
    )
    body(doc, "Ensures the system mostly (95%) picks the best outfit but occasionally (5%) picks the second-best, creating session variety. This prevents users from seeing the exact same outfit every time they generate a recommendation.")

    h2(doc, "5.2 Two-Day Garment Cooldown")
    code_block(doc,
        "recently_used = [\n"
        "    item_id for (item_id, used_day) in used_items_with_day\n"
        "    if day_idx - used_day < 2  # Used within last 2 days\n"
        "]\n"
        "# Items used on day 0 are excluded from days 0,1 but allowed on day 2+"
    )
    body(doc, "Implements a sliding window cooldown. A garment worn on Monday (day 0) is excluded from Tuesday (day 1) but can re-appear from Wednesday (day 2) onwards. This mirrors realistic wardrobe behaviour.")

    h2(doc, "5.3 Colour Harmony Scoring")
    code_block(doc,
        "def color_harmony_score(garments):\n"
        "    pairs = all combinations of 2 garments\n"
        "    for each pair:\n"
        "        if complementary:     score += 1.0  # Max contrast\n"
        "        elif triadic:         score += 0.9  # Three-way balance\n"
        "        elif analogous:       score += 0.7  # Subtle harmony\n"
        "        elif monochromatic:   score += 0.6  # Same hue, diff shade\n"
        "        elif neutral:         score += 0.8  # Neutral always safe\n"
        "    return average_score (0–1)"
    )
    body(doc, "Each garment pair is evaluated for harmony type. The average across all pairs gives the outfit's colour harmony score, which is then multiplied by 30 to yield points out of 30.")

    doc.save(os.path.join(OUT_DIR, "3_Libraries_Components_Explanation.docx"))
    print("✅ File 3 saved.")


# =============================================================================
# FILE 4: BASICS TO ADVANCED UNDERSTANDING
# =============================================================================

def build_file4():
    doc = Document()
    add_title_page(
        doc,
        "Adorkable AI — Basics to Advanced Understanding",
        "File 4 of 4 | Comprehensive Learning Guide"
    )

    # ---- PART 1: BASICS ----
    h1(doc, "PART 1: Basic Concepts")

    h2(doc, "1.1 What is Adorkable AI?")
    body(doc,
        "Adorkable AI is a web application that acts as your personal AI fashion stylist. "
        "You upload photos of your clothes, upload a selfie and a body photo, and the system "
        "automatically learns your skin tone and body shape. From that point forward, every "
        "time you ask for an outfit recommendation, the AI considers all of these factors "
        "together with the current weather in your city to suggest the best possible outfit "
        "from YOUR own wardrobe."
    )

    h2(doc, "1.2 What is a Recommendation System?")
    body(doc,
        "A recommendation system is software that predicts what a user might like based on "
        "available data. Netflix recommends movies, Spotify recommends songs, and Adorkable AI "
        "recommends outfits. Most recommendation systems use one of three approaches:"
    )
    bullet(doc, "Collaborative Filtering: 'Users like you also liked...' — requires large user datasets.")
    bullet(doc, "Content-Based Filtering: 'Based on what you own...' — uses item attributes.")
    bullet(doc, "Rule-Based Expert Systems: 'According to fashion rules...' — uses domain knowledge.")
    body(doc,
        "Adorkable AI uses a hybrid of content-based filtering (it knows your garments' colours, "
        "categories, and style) and a rule-based expert system (it knows colour harmony theory, "
        "body shape dressing rules, and weather-appropriate fabric rules)."
    )

    h2(doc, "1.3 What is Computer Vision?")
    body(doc,
        "Computer vision is the field of AI that enables computers to interpret images. "
        "In Adorkable AI, computer vision is used in two ways:"
    )
    bullet(doc, "Skin Tone Analysis: The computer 'looks at' your selfie, finds your face, and measures the colour of your cheek skin.")
    bullet(doc, "Body Shape Analysis: The computer 'looks at' your body photo, finds your skeleton (33 joint points), and measures your shoulder and hip widths to classify your body shape.")
    body(doc, "Both tasks are handled by Google's MediaPipe library, which contains pre-trained neural network models for these tasks.")

    h2(doc, "1.4 What is REST API?")
    body(doc,
        "REST (Representational State Transfer) API is a way for two software systems to "
        "communicate over the internet using standard HTTP methods (GET, POST, DELETE). "
        "In Adorkable AI, the Streamlit frontend and the FastAPI backend are separate programs "
        "that communicate via REST:"
    )
    bullet(doc, "Frontend sends: POST /api/v1/recommend with {'occasion': 'Work', 'style_pref': 'Western'}")
    bullet(doc, "Backend responds: {'score': 87.4, 'top': {...}, 'bottom': {...}, 'why_this_suits_you': '...'}")

    doc.add_page_break()

    # ---- PART 2: HOW AI FASHION RECOMMENDATION WORKS ----
    h1(doc, "PART 2: How the AI Fashion Recommendation Works")

    h2(doc, "2.1 The Complete Flow")
    body(doc, "When you click 'Get My Outfit', this is exactly what happens in order:")
    steps = [
        ("Step 1", "Frontend sends your occasion, style preference, and JWT token to POST /api/v1/recommend"),
        ("Step 2", "Backend verifies your JWT token and identifies you"),
        ("Step 3", "Backend loads your wardrobe from the database"),
        ("Step 4", "Backend fetches the current weather for your city from OpenWeatherMap API"),
        ("Step 5", "Backend loads your profile (skin tone + body shape)"),
        ("Step 6", "Engine generates all valid top+bottom and dress combinations from your wardrobe"),
        ("Step 7", "If Eastern style, only traditional_top × traditional_bottom pairs are kept"),
        ("Step 8", "Each candidate outfit is scored (colour harmony + skin + body + weather + occasion)"),
        ("Step 9", "Top-5 candidates are selected, softmax probabilities calculated, one chosen randomly"),
        ("Step 10", "If user is female, a hijab is randomly selected from wardrobe hijabs"),
        ("Step 11", "Outfit log is written to database (for analytics)"),
        ("Step 12", "JSON response sent back to frontend with outfit details and explanations"),
        ("Step 13", "Frontend displays garment images, score, why-explanation, and weather context"),
    ]
    for step, desc in steps:
        bullet(doc, f"{step}: {desc}")

    h2(doc, "2.2 The Scoring System Explained Simply")
    body(doc, "Think of the scoring system as five judges giving marks out of their maximum:")
    bullet(doc, "Judge 1 — Colour Expert (30 pts): Do the colours of the top and bottom go together? Complementary colours (opposites on colour wheel) score highest.")
    bullet(doc, "Judge 2 — Skin Tone Analyst (20 pts): Does the outfit's colours flatter your specific skin tone and undertone? Warm undertones suit earthy tones; cool undertones suit jewel tones.")
    bullet(doc, "Judge 3 — Body Shape Stylist (20 pts): Does the outfit's silhouette suit your body shape? Pear shapes are flattered by structured shoulders; inverted triangles suit wide-leg bottoms.")
    bullet(doc, "Judge 4 — Weather Reporter (20 pts): Is the fabric weight appropriate for today's temperature? Heavy wool in 35°C scores low; light cotton in -5°C scores low.")
    bullet(doc, "Judge 5 — Occasion Coordinator (10 pts): Does the outfit match the occasion? A casual t-shirt for a formal event scores low.")
    bullet(doc, "Bonus — Trend Detector (+5 pts): If your outfit matches current seasonal trends, you get a bonus.")

    doc.add_page_break()

    h2(doc, "2.3 Colour Theory Basics")
    body(doc,
        "Colour theory is a set of rules artists, designers, and fashion experts use to "
        "create visually pleasing combinations. The key concepts used in Adorkable AI:"
    )
    bullet(doc, "Colour Wheel: Colours arranged in a circle from red → orange → yellow → green → blue → violet → back to red.")
    bullet(doc, "Complementary Colours: Colours directly opposite on the wheel (blue + orange, red + green). High contrast, bold combinations.")
    bullet(doc, "Analogous Colours: Colours adjacent on the wheel (blue + blue-green + green). Harmonious, subtle combinations.")
    bullet(doc, "Monochromatic: Different shades of the same colour (navy + light blue + denim). Sophisticated, easy to execute.")
    bullet(doc, "Triadic: Three colours equally spaced on the wheel (red + blue + yellow). Vibrant, balanced.")
    body(doc,
        "The system converts all colours to HSL (Hue, Saturation, Lightness) format. "
        "The Hue is the position on the colour wheel (0°-360°). "
        "Two colours with hues 180° apart are complementary. "
        "This allows the computer to mathematically determine any colour relationship."
    )

    doc.add_page_break()

    # ---- PART 3: WARDROBE OPTIMISATION ----
    h1(doc, "PART 3: Wardrobe Optimisation Techniques")

    h2(doc, "3.1 The Capsule Wardrobe Concept")
    body(doc,
        "A capsule wardrobe is a curated collection of essential, versatile pieces that can "
        "be mixed and matched to create many outfits. The mathematical principle is: "
        "n tops × m bottoms = n×m possible outfits. With 5 tops and 4 bottoms, you have "
        "20 possible outfits. Adding one jacket multiplies this by 2 (with/without jacket). "
        "Adorkable AI's Smart Combo feature explicitly generates these combinatorial possibilities."
    )

    h2(doc, "3.2 Wear Frequency Tracking")
    body(doc,
        "Every time an outfit is recommended and the user confirms they wore it, the "
        "wear_count field increments for each garment. The analytics dashboard shows "
        "least-worn items, prompting users to rediscover forgotten pieces. This is the "
        "AI equivalent of a personal stylist who remembers what you've worn."
    )

    h2(doc, "3.3 Two-Day Rotation Algorithm")
    body(doc,
        "The weekly planner enforces a minimum 2-day gap before repeating a garment. "
        "This serves two purposes: (1) it ensures variety in daily looks, and (2) it "
        "gives garments time to air out between wears — a practical hygiene consideration "
        "built into the AI logic."
    )

    h2(doc, "3.4 Starter Wardrobe Catalog")
    body(doc,
        "For new users with empty wardrobes, the system automatically supplements their "
        "personal wardrobe with up to 10 items from a pre-curated catalog of 32 garments "
        "(gender-filtered). This ensures new users get useful recommendations immediately "
        "without needing to manually upload all their clothing first."
    )

    doc.add_page_break()

    # ---- PART 4: TECHNICAL DEEP DIVE ----
    h1(doc, "PART 4: Technical Deep Dive")

    h2(doc, "4.1 Async Programming in FastAPI")
    body(doc,
        "FastAPI uses Python's async/await syntax for non-blocking I/O. When the server "
        "makes a database query or API call, instead of blocking the entire process while "
        "waiting for a response, it yields control to handle other requests. This allows "
        "a single-process server to handle hundreds of simultaneous requests efficiently."
    )
    code_block(doc,
        "# Synchronous (blocking) — bad for servers:\nresult = database.query(...)  # All other requests wait!\n\n"
        "# Asynchronous (non-blocking) — FastAPI style:\nresult = await database.execute(select(User)...)  # Other requests can proceed"
    )
    body(doc,
        "All database operations in Adorkable AI use await to avoid blocking. "
        "AI inference (MediaPipe, KMeans) is CPU-bound and uses run_in_threadpool() "
        "to offload to a thread pool, preventing the async event loop from stalling."
    )

    h2(doc, "4.2 JWT Authentication Flow")
    body(doc, "The authentication flow follows the OAuth2 Bearer Token pattern:")
    steps = [
        ("Register/Login", "User provides email + password → backend verifies and returns JWT token"),
        ("Token Storage", "Frontend stores token in st.session_state (in-memory, not cookies)"),
        ("Protected Request", "Frontend adds 'Authorization: Bearer <token>' header to every API call"),
        ("Token Verification", "Backend's get_current_user dependency decodes token, checks expiry, fetches user from DB"),
        ("Response", "If token valid, route handler runs; if invalid, 401 Unauthorized is returned"),
    ]
    for name, desc in steps:
        bullet(doc, f"{name}: {desc}")

    h2(doc, "4.3 KMeans Colour Extraction")
    body(doc,
        "When you upload a garment image, the colour extraction pipeline does the following:"
    )
    bullet(doc, "Load image with Pillow, resize to 150×150 pixels (speed)")
    bullet(doc, "Convert to RGB, reshape to flat array of (n_pixels × 3) shape")
    bullet(doc, "Run KMeans(n_clusters=5) to find 5 dominant colour centres")
    bullet(doc, "Count pixels closest to each centre; largest count = dominant colour")
    bullet(doc, "Convert dominant RGB to hex string (#RRGGBB)")
    bullet(doc, "Map hex to nearest fashion colour name (Navy Blue, Coral Red, etc.) using pre-defined palette")
    body(doc, "KMeans is O(n_pixels × k × iterations), which runs in ~0.1 seconds on a 150×150 image.")

    h2(doc, "4.4 MediaPipe Inference Pipeline")
    code_block(doc,
        "# Skin tone flow:\nimage = cv2.imread(selfie_path)\nmp_image = mp.Image(image_format=ImageFormat.SRGB, data=cv2.cvtColor(image, cv2.COLOR_BGR2RGB))\nresult = FACE_DETECTOR.detect(mp_image)\nbbox = result.detections[0].bounding_box\ncheek_region = image[y1:y2, x1:x2]\navg_rgb = cheek_region.mean(axis=(0,1))  # Average pixel colour\nskin_tone = classify_brightness(avg_rgb)\nundertone = classify_rg_ratio(avg_rgb)"
    )
    body(doc, "The key insight is that MediaPipe returns normalised coordinates (0.0–1.0). These must be multiplied by image width/height to get pixel coordinates before cropping.")

    doc.add_page_break()

    # ---- PART 5: FUTURE ENHANCEMENTS ----
    h1(doc, "PART 5: Future Enhancements")

    h2(doc, "5.1 Short-Term (3-6 months)")
    bullet(doc, "Custom clothing classifier: Fine-tune a MobileNetV3 model on fashion datasets (DeepFashion) to automatically detect garment category and style from uploaded images, removing manual tagging.")
    bullet(doc, "Outfit history calendar: Visual calendar showing what outfits were worn on which days.")
    bullet(doc, "Social sharing: Generate shareable outfit cards with QR code linking to the outfit details.")
    bullet(doc, "Mobile-responsive UI: Adapt Streamlit frontend for mobile browser use.")

    h2(doc, "5.2 Medium-Term (6-12 months)")
    bullet(doc, "Collaborative filtering layer: Learn from aggregated user preference data (anonymised) to discover community trends by demographics.")
    bullet(doc, "Virtual try-on: Integrate ControlNet or similar diffusion model to visualise garments on the user's body photo.")
    bullet(doc, "Shopping integration: When an outfit scores low due to missing piece types, suggest specific items to purchase from partner retailers.")
    bullet(doc, "Voice interface: Integrate speech-to-text so users can say 'Show me a casual outfit for today' and receive recommendations.")

    h2(doc, "5.3 Long-Term (12+ months)")
    bullet(doc, "Multimodal LLM integration: Use a vision-language model (GPT-4V or LLaVA) to generate richer textual explanations and engage in natural language fashion conversation.")
    bullet(doc, "Smart wardrobe IoT: RFID/NFC integration to automatically detect when garments are removed from the wardrobe and update wear counts.")
    bullet(doc, "Sustainability scoring: Add a sustainability dimension to outfit scoring, rewarding use of older garments and penalising high-environmental-impact fabrics.")
    bullet(doc, "Cross-user recommendations: 'Users with similar skin tone and body shape also like...' collaborative filtering layer.")

    doc.add_page_break()

    # ---- PART 6: CONCLUSION ----
    h1(doc, "PART 6: Conclusion")
    body(doc,
        "Adorkable AI demonstrates that AI-powered personalisation is not limited to "
        "large tech companies with massive datasets. By combining well-established computer "
        "vision models (MediaPipe), classical colour theory algorithms, physics-based weather "
        "rules, and probabilistic selection mathematics, this project delivers a genuinely "
        "useful and personalised fashion recommendation system that runs entirely on consumer "
        "hardware with no cloud dependency."
    )
    body(doc,
        "The system's culturally inclusive design—supporting both Western and Eastern "
        "styles with dedicated algorithmic enforcement—sets it apart from existing fashion "
        "technology that predominantly serves Western markets. The mandatory hijab "
        "coordination for observant female users represents a thoughtful integration of "
        "religious dress practices into AI system design, a topic of growing importance "
        "in inclusive AI research."
    )
    body(doc,
        "From a software engineering perspective, the project successfully integrates "
        "six distinct technical domains (computer vision, colour science, RESTful APIs, "
        "async databases, probabilistic algorithms, and web frontend development) into "
        "a single coherent application with clean separation of concerns, shared utility "
        "modules, and comprehensive error handling. The codebase serves as a practical "
        "reference implementation for building full-stack AI applications in Python."
    )
    body(doc,
        "Future work will focus on replacing rule-based components with learned models "
        "where sufficient training data becomes available, expanding cultural fashion "
        "coverage, and developing the mobile and social features that would make Adorkable "
        "AI viable as a consumer product. The foundation built in this project provides "
        "a solid technical base for all such extensions."
    )

    doc.save(os.path.join(OUT_DIR, "4_Basics_to_Advanced_Understanding.docx"))
    print("✅ File 4 saved.")


# =============================================================================
# MAIN
# =============================================================================

if __name__ == "__main__":
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("Installing python-docx...")
        import subprocess, sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

    print("Generating documentation files...")
    build_file1()
    build_file2()
    build_file3()
    build_file4()
    print(f"\n✅ All 4 files saved to: {OUT_DIR}")
    print("Files:")
    for f in ["1_Project_Abstract_Details.docx",
              "2_Code_Line_by_Line_Explanation.docx",
              "3_Libraries_Components_Explanation.docx",
              "4_Basics_to_Advanced_Understanding.docx"]:
        path = os.path.join(OUT_DIR, f)
        size = os.path.getsize(path) // 1024
        print(f"  {f}  ({size} KB)")
