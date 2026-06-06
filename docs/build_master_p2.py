"""Part 2: Part C — Every File Explained"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1A,0x1A,0x5E); BLUE=RGBColor(0x00,0x5F,0xA3); TEAL=RGBColor(0x00,0x7A,0x7A)
DARK=RGBColor(0x22,0x22,0x22); GREEN=RGBColor(0x19,0x6F,0x3D)

def sp(p,f="F4F4F4"):
    pp=p._p.get_or_add_pPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),f); pp.append(s)

def h1(d,t):
    p=d.add_heading(t,1); r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True

def h2(d,t):
    p=d.add_heading(t,2); r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=BLUE; r.font.size=Pt(14); r.bold=True; p.paragraph_format.space_before=Pt(12)

def h3(d,t):
    p=d.add_heading(t,3); r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=TEAL; r.font.size=Pt(12); r.bold=True

def h4(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=GREEN
    p.paragraph_format.space_before=Pt(8)

def body(d,t):
    p=d.add_paragraph(t)
    for r in p.runs: r.font.size=Pt(11); r.font.color.rgb=DARK
    p.paragraph_format.space_after=Pt(4)

def bul(d,t,lv=0):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t)
    r.font.size=Pt(11); r.font.color.rgb=DARK; p.paragraph_format.left_indent=Inches(0.3+lv*0.3)

def code(d,t):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
    r=p.add_run(t); r.font.name="Courier New"; r.font.size=Pt(9); r.font.color.rgb=DARK; sp(p)

def info(d,t,f="FFF8E1"):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
    r=p.add_run(t); r.font.size=Pt(10); r.italic=True; r.font.color.rgb=RGBColor(0x5D,0x40,0x00); sp(p,f)

def lv(d,l,v):
    p=d.add_paragraph(); r1=p.add_run(f"{l}: "); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=NAVY
    r2=p.add_run(v); r2.font.size=Pt(11); r2.font.color.rgb=DARK

def tbl(d,headers,rows):
    t=d.add_table(rows=1+len(rows),cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; r=c.paragraphs[0].runs[0]
        r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        tc=c._tc; tcp=tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
        s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),"1A1A5E"); tcp.append(s)
    for ri,rd in enumerate(rows):
        row=t.rows[ri+1].cells
        for ci,ct in enumerate(rd):
            row[ci].text=str(ct)
            for pr in row[ci].paragraphs:
                for rn in pr.runs: rn.font.size=Pt(10)
        fill="F0F4FF" if ri%2==0 else "FFFFFF"
        for ci in range(len(headers)):
            tc=t.rows[ri+1].cells[ci]._tc; tcp=tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
            s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),fill); tcp.append(s)
    d.add_paragraph()

def pb(d): d.add_page_break()


def build_part_c(doc):
    h1(doc,"PART C — EVERY FILE EXPLAINED IN DETAIL")
    info(doc,"Each section: PURPOSE → KEY FUNCTIONS in plain English → CODE EXAMPLES","E3F2FD")

    # C.1
    h2(doc,"C.1  backend/main.py — Application Entry Point")
    info(doc,"FIRST file that runs when you start the server. Like a power switch that activates everything.","FFF3E0")
    lv(doc,"Purpose","Creates FastAPI app, connects all routers, configures CORS, initialises database")
    h4(doc,"lifespan(app) — Startup/shutdown handler")
    body(doc,"Runs create_all_tables() before the server accepts any requests. Creates the uploads/ folder. The 'yield' separates startup (before) from shutdown (after).")
    code(doc,"@asynccontextmanager\nasync def lifespan(app):\n    await create_all_tables()\n    os.makedirs(UPLOAD_DIR, exist_ok=True)\n    yield  # server runs here")
    h4(doc,"CORS Middleware")
    body(doc,"Allows the Streamlit frontend (port 8501) to talk to the backend (port 8006). Browsers block cross-origin requests by default — this explicitly allows them.")
    h4(doc,"app.include_router(...) — Connect URL groups")
    body(doc,"Connects all 7 router files to the main app with /api/v1 prefix. So /recommend becomes /api/v1/recommend.")
    pb(doc)

    # C.2
    h2(doc,"C.2  backend/config.py — All Settings and Constants")
    info(doc,"Change a value here → it changes everywhere. The single source of truth for all AI settings.","E8F5E9")
    tbl(doc,["Constant","Value","Controls"],[
        ["SECRET_KEY","From .env","Signs JWT tokens — must be secret"],
        ["OPENWEATHER_API_KEY","From .env","Live weather API access"],
        ["DATABASE_URL","sqlite:///./adorkable.db","Database file location"],
        ["ACCESS_TOKEN_EXPIRE_MINUTES","10080 (7 days)","How long users stay logged in"],
        ["COLOR_EXTRACTION_CLUSTERS","5","KMeans colour groups"],
        ["TOP_N_STOCHASTIC","5","Candidates considered in random pick"],
        ["BODY_SHAPE_RULES","Dict of dicts","Best silhouettes per body shape"],
        ["SCORING_WEIGHTS","Dict of floats","Points per dimension: 30,20,20,20,10,5"],
    ])
    h4(doc,"BODY_SHAPE_RULES — Most important constant")
    code(doc,'BODY_SHAPE_RULES = {\n  "Hourglass": {"best_silhouettes":["fitted","belted","wrap"],"avoid":["boxy"]},\n  "Pear":      {"best_silhouettes":["a-line","wide-leg"],"avoid":["pencil"]},\n  # 4 more shapes...\n}')
    pb(doc)

    # C.3
    h2(doc,"C.3  backend/database.py — Database Tables and Queries")
    info(doc,"Defines all 4 database tables and query helper functions.","FFF3E0")
    h4(doc,"class User — Account table")
    code(doc,"id, email (unique), hashed_password (bcrypt), gender, city, created_at")
    h4(doc,"class UserProfile — AI analysis results")
    code(doc,"id, user_id→User, skin_tone (6 levels), skin_undertone (3), body_shape (6), selfie_path, body_photo_path")
    h4(doc,"class GarmentItem — Wardrobe table")
    code(doc,'id, user_id→User, image_path, category, style, dominant_color,\ncolor_hex, fabric_weight, occasion_tags (JSON), gender_fit, wear_count, last_worn')
    h4(doc,"class OutfitLog — History table")
    code(doc,"id, user_id, top_id, bottom_id, dress_id, outerwear_id,\noccasion, score (0-105), trending_badge, worn_date, created_at")
    h4(doc,"get_db() — Session provider")
    body(doc,"Creates a fresh database connection for every API request. Automatically commits or rolls back, then closes. Prevents connection leaks.")
    h4(doc,"create_all_tables() — One-time setup")
    body(doc,"Creates all 4 tables in adorkable.db if they don't exist. Safe to call multiple times.")
    h4(doc,"get_user_wardrobe(db, user_id) → List[GarmentItem]")
    body(doc,"Queries garment_items table for all rows belonging to this user. Returns a Python list used by all recommendation engines.")
    pb(doc)

    # C.4
    h2(doc,"C.4  backend/auth.py — Complete Authentication System")
    info(doc,"Handles account creation, login, JWT token generation, and identity verification on every protected API call.","FFF8E1")
    h4(doc,"hash_password(plain) → bcrypt_hash")
    body(doc,"Scrambles the password using bcrypt. One-way — cannot be reversed. Random salt means two identical passwords produce different hashes.")
    code(doc,'hash_password("mypass") → "$2b$12$xwANv..."  (irreversible)')
    h4(doc,"verify_password(plain, hashed) → True/False")
    body(doc,"At login: re-applies bcrypt with the stored salt and compares. Returns True only if they match.")
    h4(doc,"create_access_token(data) → jwt_string")
    body(doc,"Creates a JWT: HEADER.PAYLOAD.SIGNATURE. Payload contains user ID + expiry (7 days). Signed with SECRET_KEY — cannot be forged.")
    h4(doc,"get_current_user(credentials, db) → User")
    body(doc,"FastAPI dependency — runs automatically before every protected endpoint. Extracts Bearer token, decodes it, looks up user in database, returns User object. Returns 401 if anything fails.")
    h4(doc,"POST /auth/register")
    body(doc,"Check email unique → hash password → save User → return JWT token (user logged in immediately).")
    h4(doc,"POST /auth/login")
    body(doc,"Find user by email → verify password → return JWT token.")
    pb(doc)

    # C.5
    h2(doc,"C.5  backend/engine/outfit_scorer.py — The AI Brain (Core Scorer)")
    info(doc,"THE most important file. Every recommendation, plan, and combo scores through here. Judges outfits across 5 dimensions.","FFE0B2")
    lv(doc,"Size","~598 lines")
    h4(doc,"score_outfit(garments, profile, weather, occasion, style_pref) → dict")
    body(doc,"Master function. Calls all 5 sub-scorers, adds trending bonus, returns total score out of 105 plus explanation text.")
    code(doc,"total = (colour_harmony×30) + (skin_flattery×20) + (body_shape×20)\n      + (weather×20) + (occasion×10) + (trending_bonus: 0 or 5)")
    h4(doc,"score_color_harmony(garments) → (0.0-1.0, label)")
    body(doc,"Converts hex colours to HSL. Measures hue angle between every pair. Classifies: Complementary(~180°)=1.0, Triadic(~120°)=0.9, Neutral pair=0.8, Analogous(<30°)=0.7, Monochromatic=0.6, None=0.2.")
    h4(doc,"score_skin_tone_flattery(garments, skin_tone, undertone) → (0.0-1.0, text)")
    body(doc,"Checks each garment's dominant_color against the flattering palette for the user's skin undertone (from skin_tone_palettes.json). Highly flattering=1.0, neutral=0.5, unflattering=0.1.")
    h4(doc,"score_body_shape_suitability(garments, body_shape) → (0.0-1.0, text)")
    body(doc,"Checks garment style/category against BODY_SHAPE_RULES. Best silhouette match=1.0, avoid list=0.2, neutral=0.5.")
    h4(doc,"weather_suitability_score(garment, temp_c, condition) → float")
    body(doc,"Compares garment fabric_weight to ideal weight for the temperature. Exact match=1.0, off by 1 level=0.6, off by 2=0.2. Rain+no outerwear → penalty.")
    h4(doc,"score_occasion_match(garments, occasion) → float")
    body(doc,"Parses garment occasion_tags JSON. Full match=1.0, partial=0.5, none=0.0.")
    h4(doc,"check_trending(garments) → (bool, message)")
    body(doc,"Compares garment colours/styles to trends.json for current season. Match → True + '🔥 on-trend' message.")
    pb(doc)

    # C.6
    h2(doc,"C.6  backend/engine/color_theory.py — Colour Harmony Maths")
    info(doc,"Implements colour wheel mathematics. All colour harmony checks in the scorer use this file.","E8EAF6")
    h4(doc,"hex_to_hsl(hex_string) → (hue 0-360, saturation 0-100, lightness 0-100)")
    body(doc,"Converts a hex colour code to HSL. Hue=position on colour wheel, Saturation=vivid vs grey, Lightness=bright vs dark. Required for harmony angle calculations.")
    h4(doc,"get_hue_distance(h1, h2) → degrees")
    code(doc,"diff = abs(h1-h2)\nreturn min(diff, 360-diff)  # shorter arc on the wheel")
    h4(doc,"classify_harmony(hue1, hue2) → type_string")
    body(doc,"Returns: 'Complementary'(~180°), 'Triadic'(~120°), 'Analogous'(<30°), 'Monochromatic'(<15°), 'Split-Complementary', or 'None'.")
    h4(doc,"is_neutral(color_name) → bool")
    body(doc,"Returns True for Black, White, Beige, Grey, Cream, Ivory, Tan, Camel. Neutrals pair well with any colour.")
    pb(doc)

    # C.7
    h2(doc,"C.7  backend/engine/weather_rules.py — Temperature & Fabric Rules")
    info(doc,"All fashion rules for weather appropriateness. Knows which fabrics suit which temperatures.","E0F2F1")
    h4(doc,"get_fabric_weight_needed(temp_c) → string")
    code(doc,"temp >= 28 → 'light'  (silk, chiffon)\n22-28     → 'light-medium'  (cotton, rayon)\n15-22     → 'medium'  (denim, jersey)\n10-15     → 'medium-heavy'  (light wool)\ntemp < 10 → 'heavy'  (wool, tweed, fleece)")
    h4(doc,"weather_suitability_score(garment, temp_c, condition, has_outerwear) → float")
    body(doc,"Distance 0=1.0, distance 1=0.6, distance 2=0.2, distance 3+=0.0. Rain penalty -0.3 if no outerwear.")
    h4(doc,"should_include_outerwear(weather, has_outerwear) → bool")
    body(doc,"Returns True when temp<12°C OR condition is Rain/Thunderstorm. Forces outerwear into outfit candidate generation.")
    h4(doc,"generate_weather_explanation(outfit_garments, weather) → string")
    body(doc,"Returns the human-readable weather context shown on the outfit card, e.g. 'It's 28°C and sunny — breathable cotton is ideal today.'")
    pb(doc)

    # C.8
    h2(doc,"C.8  backend/engine/outfit_constraints.py — Pairing Validity Rules")
    info(doc,"Decides which garment combinations are VALID before they are scored. Acts as a filter gate.","FFF3E0")
    h4(doc,"outfit_has_complete_separates_or_dress(scored_dict) → bool")
    code(doc,"Valid: (top AND bottom) OR (dress alone)\nInvalid: dress + top, dress + bottom, top alone, bottom alone")
    h4(doc,"generate_wardrobe_outfit_candidates(wardrobe, profile, weather, occasion, style_pref) → List")
    body(doc,"Loops through all tops×bottoms and all dresses. For Eastern style: skips any pair where both are not 'traditional_*' category. Scores each valid combination. Returns only outfits above min_score threshold.")
    code(doc,"for top in tops:\n  for bottom in bottoms:\n    if style=='Eastern' and not both_traditional: continue\n    scored = score_outfit([top,bottom],...)\n    if scored['score'] >= min_score: candidates.append(scored)")
    h4(doc,"normalize_starter_gender(gender_string) → 'male'/'female'/'unisex'")
    body(doc,"Converts 'F','Woman','FEMALE','female' → 'female'. Converts 'M','Man','MALE' → 'male'. Used to filter starter catalog by user's gender.")
    pb(doc)

    # C.9
    h2(doc,"C.9  backend/engine/stochastic_selector.py — Smart Random Selection")
    info(doc,"Ensures variety: picks the best outfit most often but not always the same one.","E8F5E9")
    h4(doc,"softmax(weights) → probabilities list")
    code(doc,"scores=[85,82,78]\nmax_w=85\nexp_w=[e^0, e^-3, e^-7] = [1.0, 0.05, 0.001]\nprobs=[0.951, 0.047, 0.001]  # sum=1.0")
    h4(doc,"weighted_random_select(candidates, top_n=5) → one_outfit")
    body(doc,"Sorts by score, takes top 5, applies softmax, uses random.choices() with those probabilities. Best outfit chosen ~95% of the time, others rarely.")
    h4(doc,"select_with_exclusion(candidates, excluded_ids) → outfit or None")
    body(doc,"Filters out candidates using any garment in excluded_ids (the 2-day cooldown list). Returns None if nothing remains — triggers fallback in the planner.")
    pb(doc)

    # C.10
    h2(doc,"C.10  backend/engine/weekly_planner.py — 7-Day Planner")
    info(doc,"Generates 7 outfits with 2-day garment cooldown and bottom variety boosting.","FFF8E1")
    h4(doc,"generate_weekly_plan(profile, wardrobe, city, occasions, style_pref) → List[7]")
    body(doc,"Loops day 0-6. Each day: fetch weather → generate candidates → block recently-used garments → boost unused bottoms (+15pts) → select_with_exclusion → record used garments.")
    h4(doc,"_boost_unused_bottoms(candidates, used_bottom_ids) → candidates")
    body(doc,"Adds 15.0 to score of outfits using bottoms NOT in the used list. Strongly encourages different trousers/skirts each day.")
    h4(doc,"generate_quick_plan(wardrobe, profile, city, style_pref) → List[7]")
    body(doc,"Uses default occasions: Casual, Academic, Casual, Formal, Casual, Casual, Formal.")
    h4(doc,"get_weekly_stats(weekly_plan) → dict")
    body(doc,"Returns: average_score, trending_days count, unique_garments_used count.")
    pb(doc)

    # C.11
    h2(doc,"C.11  backend/engine/combo_generator.py — Smart Combo Generator")
    info(doc,"Given one garment the user picks, finds all outfits that work well with it.","E0F2F1")
    h4(doc,"generate_combos(selected_item, all_wardrobe, ...) → List[dicts]")
    body(doc,"Behaviour by category:")
    code(doc,"TOP selected    → pair with all compatible BOTTOMS\nBOTTOM selected → pair with all compatible TOPS\nDRESS selected  → score alone (+ outerwear if cold)\nOUTERWEAR       → try with every top+bottom pair and every dress")
    body(doc,"Eastern style check applies: only traditional_top+traditional_bottom when Eastern preference set.")
    h4(doc,"find_alternatives(outfit, all_wardrobe, ...) → dict")
    body(doc,"For each piece in an outfit, tests all garments of the same category in that slot. Returns top 3 alternatives per position with their scores. Powers the 'swap this piece' feature.")
    pb(doc)

    # C.12
    h2(doc,"C.12  backend/ml/skin_tone.py — Skin Tone Analysis")
    info(doc,"Uses BlazeFace TFLite model to detect face, samples cheek pixels, classifies skin tone.","FFE0B2")
    h4(doc,"download_model_if_needed()")
    body(doc,"Checks if blaze_face_short_range.tflite exists in ml/models/. Downloads from Google MediaPipe registry if missing. Runs silently on first use.")
    h4(doc,"detect_face_bbox(image_path) → (x,y,w,h) normalised or None")
    body(doc,"Runs BlazeFace on the image. Returns the face bounding box as 0.0–1.0 normalised coordinates. Returns None if no face found.")
    h4(doc,"extract_cheek_region(image_array, bbox) → numpy_array")
    body(doc,"Converts normalised bbox to pixels. Crops left cheek area (left third, middle vertical third of face). Avoids nose (shadow) and forehead (highlight) which distort colour.")
    h4(doc,"classify_skin_tone(avg_rgb) → (tone_label, undertone_label)")
    code(doc,"brightness=(R+G+B)/3\n>220→Very Fair, 185-220→Fair, 155-185→Light Medium\n125-155→Medium, 90-125→Medium Dark, <90→Dark\n\nratio=R/G\n>1.15→Warm, <0.95→Cool, else→Neutral")
    h4(doc,"analyse_skin_tone(image_path) → result_dict")
    body(doc,"Complete pipeline: detect face → crop cheek → average RGB → classify. Returns {skin_tone, skin_undertone, confidence} or error dict.")
    pb(doc)

    # C.13
    h2(doc,"C.13  backend/ml/body_shape.py — Body Shape Classification")
    info(doc,"Uses PoseLandmarker TFLite to find 33 body joints and classify shape from shoulder/hip ratio.","E8EAF6")
    h4(doc,"detect_landmarks(image_path) → List[33 landmarks] or None")
    body(doc,"Runs PoseLandmarker. Returns 33 points each with .x, .y (normalised 0-1) and .visibility (confidence). Returns None if no person found.")
    h4(doc,"calculate_body_measurements(landmarks, image_width) → dict")
    code(doc,"shoulder_width = pixel_dist(landmark[11], landmark[12], img_w)\nhip_width      = pixel_dist(landmark[23], landmark[24], img_w)\nratio          = shoulder_width / hip_width")
    h4(doc,"classify_body_shape(measurements) → shape_string")
    code(doc,"ratio>1.15  → Inverted Triangle\nratio<0.85  → Pear\n≈1.0+narrow waist → Hourglass\n≈1.0+uniform → Rectangle\nwaist>shoulders → Apple\nelse → Athletic")
    h4(doc,"analyse_body_shape(image_path) → result_dict")
    body(doc,"Pipeline: detect landmarks → calculate measurements → classify shape. Returns {body_shape, shoulder_width, hip_width, ratio}.")
    pb(doc)

    # C.14
    h2(doc,"C.14  backend/ml/color_extractor.py — Garment Colour Extraction")
    info(doc,"Automatically finds the dominant colour of any garment photo using KMeans clustering.","FFF8E1")
    h4(doc,"extract_dominant_color(image_path, n_clusters=5) → (hex, color_name)")
    code(doc,"1. Load image with OpenCV\n2. Convert BGR→RGB\n3. Resize to 150×150 (22,500 pixels)\n4. Reshape to (22500, 3) array\n5. KMeans(n_clusters=5).fit(pixels)\n6. Find cluster with most pixels → dominant\n7. Centroid RGB → hex → nearest colour name")
    h4(doc,"hex_to_color_name(hex_string) → fashion_name")
    body(doc,"Loads color_mapping.json (100+ fashion colours). For each mapping entry, computes Euclidean RGB distance. Returns the name of the closest match (e.g. '#000080' → 'Navy Blue').")
    pb(doc)

    # C.15
    h2(doc,"C.15  backend/ml/classifier.py — Garment Attribute Classifier")
    info(doc,"Rule-based classifier (not a neural network) that assigns category, style, and attributes to uploaded garments.","E0F2F1")
    h4(doc,"classify_garment(category, style, occasions, fabric_weight) → dict")
    body(doc,"Takes user-provided metadata from the upload form. Validates against allowed values. Assigns defaults for missing fields. Returns standardised attribute dictionary saved to the database.")
    h4(doc,"infer_fabric_weight_from_category(category, style) → string")
    body(doc,"Rule-based defaults: dress/blouse → 'light', jeans/trousers → 'medium', outerwear → 'heavy', traditional_top → 'medium'.")
    pb(doc)

    # C.16
    h2(doc,"C.16  backend/routers/wardrobe.py — Wardrobe Management API")
    info(doc,"All endpoints for uploading, viewing, editing, and deleting clothes.","FFF3E0")
    h4(doc,"POST /api/v1/wardrobe/upload")
    body(doc,"Steps: validate image (Pillow) → save to disk (UUID filename) → extract_dominant_color() → classify_garment() → save GarmentItem to database. Returns garment details with assigned colour.")
    h4(doc,"GET /api/v1/wardrobe/")
    body(doc,"Returns full list of all GarmentItem rows for the current user. Powers the wardrobe grid display.")
    h4(doc,"DELETE /api/v1/wardrobe/{garment_id}")
    body(doc,"Verifies ownership → delete image file from disk → delete GarmentItem from database.")
    h4(doc,"GET /api/v1/wardrobe/preloaded")
    body(doc,"Returns starter catalog items filtered by user's gender. Shown as importable starter wardrobe.")
    h4(doc,"POST /api/v1/wardrobe/import-starter")
    body(doc,"Imports selected starter catalog items: copies image files + creates GarmentItem rows in database.")
    pb(doc)

    # C.17
    h2(doc,"C.17  backend/routers/recommendations.py — Daily Outfit Endpoint")
    info(doc,"The most-used endpoint: generates a single personalised daily outfit.","FFE0B2")
    h4(doc,"POST /api/v1/recommend")
    body(doc,"Request: {occasion, style_pref, reimagine_step}")
    code(doc,"1. Verify JWT → get user\n2. Load wardrobe (+ starter if empty)\n3. Fetch weather for user's city\n4. Load user_profile (skin+body)\n5. generate_wardrobe_outfit_candidates()\n6. weighted_random_select(skip=reimagine_step)\n7. Add hijab if female\n8. Save OutfitLog + increment wear_count\n9. Return JSON with outfit + explanations")
    body(doc,"reimagine_step: increments when user clicks 'Re-imagine'. Skips that many top candidates to cycle through different options.")
    pb(doc)

    # C.18-C.25
    h2(doc,"C.18  backend/routers/planner.py — Weekly Planner Endpoints")
    h4(doc,"POST /api/v1/plan/weekly"); body(doc,"Accepts 7 occasion strings. Calls generate_weekly_plan(). Returns 7 outfit dicts with garments, scores, day names, dates, weather.")
    h4(doc,"GET /api/v1/plan/quick"); body(doc,"No input needed. Calls generate_quick_plan() with default occasions. Returns 7-day plan instantly.")

    h2(doc,"C.19  backend/routers/combo.py — Smart Combo Endpoint")
    h4(doc,"POST /api/v1/combo/generate"); body(doc,"User sends one garment ID. Calls generate_combos(selected_garment, ...). Returns all compatible outfit combinations sorted by score.")

    h2(doc,"C.20  backend/routers/analytics.py — Analytics Endpoints")
    h4(doc,"GET /analytics/wardrobe-colors"); body(doc,"Counts garments by colour → used by donut chart.")
    h4(doc,"GET /analytics/garment-usage"); body(doc,"Garments sorted by wear_count → used by bar chart.")
    h4(doc,"GET /analytics/combinability"); body(doc,"Calculates: (n_tops × n_bottoms + n_dresses) × (n_outerwear+1). Shows how many unique outfits are theoretically possible.")
    h4(doc,"GET /analytics/outfit-history"); body(doc,"Last 30 outfit logs → used by score trend line chart.")
    h4(doc,"GET /analytics/dashboard-summary"); body(doc,"Returns: total items, most-worn, least-worn, favourite colour, last outfit score.")

    h2(doc,"C.21  backend/routers/profile.py — Profile Endpoints")
    h4(doc,"POST /profile/selfie"); body(doc,"Saves selfie → runs analyse_skin_tone() → saves results to user_profiles table.")
    h4(doc,"POST /profile/body-photo"); body(doc,"Saves body photo → runs analyse_body_shape() → saves body_shape to user_profiles.")
    h4(doc,"GET /profile/"); body(doc,"Returns current profile: skin_tone, skin_undertone, body_shape, photo paths.")
    pb(doc)

    h2(doc,"C.22  backend/routers/helpers.py — Shared Utilities")
    info(doc,"Prevents code duplication across recommendations, planner, and combo routers.","E8F5E9")
    h4(doc,"select_hijab_for_female(gender, items, day_idx, rotate) → GarmentItem or None")
    body(doc,"Returns None if user is not female. Finds all hijab items. If rotate=True (weekly planner): hijabs[day_idx % count] — ensures different hijab each day. If rotate=False (daily): random.choice — ensures session variety.")
    h4(doc,"garment_to_dict(garment) → dict or None")
    body(doc,"Converts GarmentItem object to JSON-serialisable dict: {id, category, dominant_color, color_hex, image_path}.")
    pb(doc)

    h2(doc,"C.23  backend/services/starter_wardrobe.py — Default Wardrobe")
    h4(doc,"ensure_user_has_wardrobe(db, user_id) → List[GarmentItem]")
    body(doc,"If user wardrobe is empty: loads preloaded_wardrobe.json, filters by gender, copies images to uploads folder, saves GarmentItem rows to database. New users can use the app immediately without uploading anything.")

    h2(doc,"C.24  backend/utils/weather_api.py — Weather Client")
    h4(doc,"get_current_weather(city) → dict")
    body(doc,"Calls OpenWeatherMap current API. Returns {temp_c, condition, humidity}. Falls back to {22°C, Clear, 50%} if API key missing or call fails.")
    h4(doc,"get_7day_forecast(city) → List[7 dicts]")
    body(doc,"Calls 5-day/3-hour forecast API (40 data points). Groups by date, averages temperature, finds most common condition per day. Returns up to 7 daily forecasts.")
    pb(doc)

    h2(doc,"C.25  backend/utils/image_utils.py — Image File Utilities")
    h4(doc,"save_uploaded_image(file_bytes, user_id, filename) → path")
    body(doc,"Saves image to uploads/<uuid>.ext using UUID to prevent filename collisions.")
    h4(doc,"is_valid_image(file_bytes) → bool")
    body(doc,"Attempts Pillow open — returns False for corrupt files or non-images.")
    h4(doc,"delete_image_file(path)"); body(doc,"Deletes file from disk. No error if already missing.")
    pb(doc)

    # Frontend pages
    h2(doc,"C.26  frontend/app.py — Landing Page")
    body(doc,"Home page at http://localhost:8501. Shows brand name and navigation. Checks st.session_state.token — if logged in, shows dashboard links; if not, shows Register/Login button.")
    body(doc,"Key concept: st.session_state is a dictionary that persists between button clicks, storing the JWT token, user_id, email, and gender.")

    h2(doc,"C.27  frontend/pages/1_Register_Login.py — Auth Page")
    body(doc,"Two-tab layout. Register tab: collects email, password, gender, city → POST /auth/register → stores JWT in session_state. Login tab: email + password → POST /auth/login → stores JWT. Auto-redirects to Profile page after success.")

    h2(doc,"C.28  frontend/pages/2_My_Profile.py — Profile Page")
    body(doc,"Three sections: 1) Current profile display (skin tone + body shape badges). 2) Selfie upload → POST /profile/selfie → shows AI-detected skin tone + colour palette. 3) Body photo upload → POST /profile/body-photo → shows detected body shape + style tips.")

    h2(doc,"C.29  frontend/pages/3_My_Wardrobe.py — Wardrobe Page")
    body(doc,"Three tabs: My Clothes (grid of all garments with images, colour badges, wear counts). Upload New (form for uploading garment + metadata). Starter Wardrobe (grid of preloaded items to import).")

    h2(doc,"C.30  frontend/pages/4_Daily_Outfit.py — Daily Outfit Page")
    body(doc,"User selects occasion + style. Clicks 'Get My Outfit'. Outfit appears as garment image cards with score bar, colour harmony badge, weather explanation, 'why this suits you' text. 'Re-imagine' button cycles to different outfit.")

    h2(doc,"C.31  frontend/pages/5_Weekly_Planner.py — Weekly Planner Page")
    body(doc,"User selects style preference. Clicks 'Plan My Week'. Displays 7 day-cards (Monday–Sunday) each showing garment images, score, weather, and occasion. Different garments appear each day due to cooldown rules.")

    h2(doc,"C.32  frontend/pages/6_Smart_Combo.py — Smart Combo Page")
    body(doc,"User clicks any garment from their wardrobe. All compatible outfit combinations appear sorted by score. Each combo shows paired garments with their combined score and explanation.")

    h2(doc,"C.33  frontend/pages/7_Analytics.py — Analytics Page")
    body(doc,"Dashboard with 4 interactive Plotly charts: colour distribution donut, wear frequency bar chart, outfit score trend line, combinability metric. Shows statistics cards: total items, most/least worn, favourite colour.")

    pb(doc)
    h2(doc,"C.34  backend/data/ — JSON Knowledge Files")
    tbl(doc,["File","Contents","Used By"],[
        ["preloaded_wardrobe.json","32 starter garments with full metadata","starter_wardrobe.py, wardrobe.py"],
        ["trends.json","2026 Spring/Summer + Autumn/Winter trending colours and styles","outfit_scorer.py (trending bonus)"],
        ["color_mapping.json","100+ fashion colour names with hex codes and families","color_extractor.py (name lookup)"],
        ["skin_tone_palettes.json","Flattering colour palettes per skin tone + undertone","outfit_scorer.py (skin flattery)"],
    ])
    pb(doc)


if __name__=="__main__":
    doc=Document()
    build_part_c(doc)
    out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"_master_p2.docx")
    doc.save(out)
    print(f"Part 2 saved → {out}")
