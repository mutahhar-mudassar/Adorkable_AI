"""
Adorkable AI Master Doc — Part 1: Helpers + Title + TOC + Part A + Part B
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x1A,0x1A,0x5E)
BLUE   = RGBColor(0x00,0x5F,0xA3)
TEAL   = RGBColor(0x00,0x7A,0x7A)
DARK   = RGBColor(0x22,0x22,0x22)
GREY   = RGBColor(0x55,0x55,0x55)
GREEN  = RGBColor(0x19,0x6F,0x3D)
ORANGE = RGBColor(0xCA,0x6F,0x1E)

def shade_p(p, fill="F4F4F4"):
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),fill)
    pPr.append(s)

def h1(doc,t,c=None):
    c=c or NAVY
    p=doc.add_heading(t,level=1)
    r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=c; r.font.size=Pt(18); r.bold=True
    p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(6)
    return p

def h2(doc,t,c=None):
    c=c or BLUE
    p=doc.add_heading(t,level=2)
    r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=c; r.font.size=Pt(14); r.bold=True
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
    return p

def h3(doc,t,c=None):
    c=c or TEAL
    p=doc.add_heading(t,level=3)
    r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=c; r.font.size=Pt(12); r.bold=True
    return p

def h4(doc,t):
    p=doc.add_paragraph()
    r=p.add_run(t); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=GREEN
    p.paragraph_format.space_before=Pt(8)
    return p

def body(doc,t,italic=False):
    p=doc.add_paragraph(t)
    for r in p.runs:
        r.font.size=Pt(11); r.font.color.rgb=DARK
        if italic: r.italic=True
    p.paragraph_format.space_after=Pt(4)
    return p

def bullet(doc,t,level=0):
    p=doc.add_paragraph(style="List Bullet")
    r=p.add_run(t); r.font.size=Pt(11); r.font.color.rgb=DARK
    p.paragraph_format.left_indent=Inches(0.3+level*0.3)
    return p

def code(doc,t):
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Inches(0.3)
    p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
    r=p.add_run(t); r.font.name="Courier New"; r.font.size=Pt(9)
    r.font.color.rgb=RGBColor(0x10,0x10,0x10)
    shade_p(p,"F4F4F4"); return p

def info(doc,t,fill="FFF8E1"):
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Inches(0.3); p.paragraph_format.right_indent=Inches(0.3)
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    r=p.add_run(t); r.font.size=Pt(10); r.italic=True
    r.font.color.rgb=RGBColor(0x5D,0x40,0x00)
    shade_p(p,fill); return p

def lv(doc,label,val):
    p=doc.add_paragraph()
    r1=p.add_run(f"{label}: "); r1.bold=True; r1.font.size=Pt(11); r1.font.color.rgb=NAVY
    r2=p.add_run(val); r2.font.size=Pt(11); r2.font.color.rgb=DARK
    p.paragraph_format.space_after=Pt(3); return p

def table(doc,headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers))
    t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        r=c.paragraphs[0].runs[0]; r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        tc=c._tc; tcp=tc.get_or_add_tcPr()
        s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear")
        s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),"1A1A5E"); tcp.append(s)
    for ri,rd in enumerate(rows):
        row=t.rows[ri+1].cells
        for ci,ct in enumerate(rd):
            row[ci].text=str(ct)
            for pr in row[ci].paragraphs:
                for rn in pr.runs: rn.font.size=Pt(10)
        fill="F0F4FF" if ri%2==0 else "FFFFFF"
        for ci in range(len(headers)):
            tc=t.rows[ri+1].cells[ci]._tc; tcp=tc.get_or_add_tcPr()
            s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear")
            s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),fill); tcp.append(s)
    doc.add_paragraph(); return t

def pb(doc): doc.add_page_break()


def build_title(doc):
    for _ in range(4): doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=t.add_run("ADORKABLE AI"); r.bold=True; r.font.size=Pt(36); r.font.color.rgb=NAVY
    s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=s.add_run("AI-Powered Personal Fashion Stylist"); r2.font.size=Pt(18)
    r2.font.color.rgb=BLUE; r2.italic=True
    doc.add_paragraph()
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=sub.add_run("MASTER PROJECT GUIDE — COMPLETE EXPLANATION")
    r3.bold=True; r3.font.size=Pt(20); r3.font.color.rgb=TEAL
    doc.add_paragraph()
    d=doc.add_paragraph(); d.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r4=d.add_run(
        "Every file, every function, every concept explained in plain English.\n"
        "No programming knowledge required to understand this document.\n\n"
        "Academic Year 2025–2026\n"
        "Department of Computer Science & Artificial Intelligence"
    )
    r4.font.size=Pt(12); r4.font.color.rgb=GREY
    pb(doc)


def build_toc(doc):
    h1(doc,"TABLE OF CONTENTS",NAVY)
    info(doc,
        "In Microsoft Word: press Ctrl+F to search for any section title. "
        "Or use View → Navigation Pane to browse all headings.", "E8F5E9")
    toc=[
        ("PART A","HOW THE WHOLE SYSTEM WORKS — START HERE"),
        ("A.1","What is Adorkable AI? (Plain English)"),
        ("A.2","The Big Picture — System Architecture & Folder Structure"),
        ("A.3","Step-by-Step: How a Recommendation is Generated"),
        ("A.4","How the AI Detects Your Skin Tone"),
        ("A.5","How the AI Detects Your Body Shape"),
        ("A.6","How Outfit Scoring Works — All 5 Dimensions Explained"),
        ("A.7","How Variety is Guaranteed — The Softmax Algorithm"),
        ("A.8","How the Weekly Planner Works"),
        ("A.9","How Eastern Style Rules are Enforced"),
        ("A.10","How the Database Stores Everything"),
        ("PART B","TECHNOLOGY STACK — EVERY LIBRARY EXPLAINED"),
        ("B.1","Full Library Reference (16 Libraries)"),
        ("B.2","How Libraries Work Together — Complete Data Journeys"),
        ("PART C","EVERY FILE EXPLAINED ONE BY ONE"),
        ("C.1","backend/main.py"),
        ("C.2","backend/config.py"),
        ("C.3","backend/database.py"),
        ("C.4","backend/auth.py"),
        ("C.5","backend/engine/outfit_scorer.py"),
        ("C.6","backend/engine/color_theory.py"),
        ("C.7","backend/engine/weather_rules.py"),
        ("C.8","backend/engine/outfit_constraints.py"),
        ("C.9","backend/engine/stochastic_selector.py"),
        ("C.10","backend/engine/weekly_planner.py"),
        ("C.11","backend/engine/combo_generator.py"),
        ("C.12","backend/ml/skin_tone.py"),
        ("C.13","backend/ml/body_shape.py"),
        ("C.14","backend/ml/color_extractor.py"),
        ("C.15","backend/ml/classifier.py"),
        ("C.16","backend/routers/wardrobe.py"),
        ("C.17","backend/routers/recommendations.py"),
        ("C.18","backend/routers/planner.py"),
        ("C.19","backend/routers/combo.py"),
        ("C.20","backend/routers/analytics.py"),
        ("C.21","backend/routers/profile.py"),
        ("C.22","backend/routers/helpers.py"),
        ("C.23","backend/services/starter_wardrobe.py"),
        ("C.24","backend/utils/weather_api.py"),
        ("C.25","backend/utils/image_utils.py"),
        ("C.26","frontend/app.py"),
        ("C.27","frontend/pages/1_Register_Login.py"),
        ("C.28","frontend/pages/2_My_Profile.py"),
        ("C.29","frontend/pages/3_My_Wardrobe.py"),
        ("C.30","frontend/pages/4_Daily_Outfit.py"),
        ("C.31","frontend/pages/5_Weekly_Planner.py"),
        ("C.32","frontend/pages/6_Smart_Combo.py"),
        ("C.33","frontend/pages/7_Analytics.py"),
        ("C.34","backend/data/ — JSON Knowledge Files"),
        ("PART D","COMPLETE DATA FLOW WALKTHROUGH"),
        ("PART E","ALGORITHMS EXPLAINED IN PLAIN ENGLISH"),
        ("PART F","HOW TO EXPLAIN THIS PROJECT TO ANYONE"),
    ]
    table(doc,["Section","Title"],toc)
    pb(doc)


def build_part_a(doc):
    h1(doc,"PART A — HOW THE WHOLE SYSTEM WORKS",NAVY)
    info(doc,"READ THIS FIRST. Plain English explanation of the entire project before any code.","E3F2FD")

    h2(doc,"A.1  What is Adorkable AI?")
    body(doc,"Adorkable AI is a website that acts as your personal AI fashion stylist. Imagine having a friend who:")
    bullet(doc,"Knows every piece of clothing you own")
    bullet(doc,"Knows your skin tone and what colours look best on you")
    bullet(doc,"Knows your body shape and what silhouettes flatter your figure")
    bullet(doc,"Checks today's weather and suggests appropriate fabric weights")
    bullet(doc,"Remembers what you wore recently so it never repeats the same outfit")
    bullet(doc,"Plans your entire week of outfits in advance")
    bullet(doc,"Respects your cultural style — Western fashion OR Eastern/traditional clothing")
    bullet(doc,"Automatically adds a matching hijab for female users who wear one")
    body(doc,"All of this runs as a web application. You open a browser, log in, upload photos of your clothes and a selfie, and the AI does everything automatically.")

    pb(doc)
    h2(doc,"A.2  The Big Picture — System Architecture")
    info(doc,"Think of it like a restaurant: the Frontend is the dining room (what users see), the Backend is the kitchen (where AI work happens), and the API is the waiter connecting them.","FFF3E0")

    body(doc,"THREE LAYERS:")
    bullet(doc,"FRONTEND — The website (Streamlit, runs at http://localhost:8501)")
    bullet(doc,"  → 8 pages: Landing, Register/Login, Profile, Wardrobe, Daily Outfit, Weekly Planner, Smart Combo, Analytics",1)
    bullet(doc,"BACKEND — The AI server (FastAPI, runs at http://localhost:8006)")
    bullet(doc,"  → 15+ API endpoints, all AI logic, database access",1)
    bullet(doc,"DATABASE — adorkable.db (SQLite file — like a spreadsheet saved on disk)")
    bullet(doc,"  → 4 tables: users, user_profiles, garment_items, outfit_logs",1)

    body(doc,"Complete folder structure:")
    code(doc,
"adorkable_ai/\n"
"├── backend/\n"
"│   ├── main.py               ← Starts the server (entry point)\n"
"│   ├── auth.py               ← Login / Register / JWT tokens\n"
"│   ├── config.py             ← All settings and AI weights\n"
"│   ├── database.py           ← 4 database table definitions\n"
"│   ├── engine/               ← The AI scoring brain\n"
"│   │   ├── outfit_scorer.py  ← Master scorer (5 dimensions)\n"
"│   │   ├── color_theory.py   ← Colour harmony mathematics\n"
"│   │   ├── weather_rules.py  ← Temperature/fabric rules\n"
"│   │   ├── outfit_constraints.py  ← Valid pairing rules\n"
"│   │   ├── stochastic_selector.py ← Softmax random selection\n"
"│   │   ├── weekly_planner.py ← 7-day planner with cooldown\n"
"│   │   └── combo_generator.py ← Smart outfit combos\n"
"│   ├── ml/                   ← Computer Vision & ML\n"
"│   │   ├── skin_tone.py      ← Face detection + skin colour\n"
"│   │   ├── body_shape.py     ← Pose detection + body shape\n"
"│   │   ├── color_extractor.py ← KMeans colour from images\n"
"│   │   └── classifier.py     ← Garment attribute detection\n"
"│   ├── routers/              ← API URL handlers\n"
"│   │   ├── wardrobe.py       ← Upload/manage clothes\n"
"│   │   ├── recommendations.py ← Daily outfit endpoint\n"
"│   │   ├── planner.py        ← Weekly plan endpoints\n"
"│   │   ├── combo.py          ← Smart combo endpoint\n"
"│   │   ├── analytics.py      ← Charts/stats endpoints\n"
"│   │   ├── profile.py        ← Profile endpoints\n"
"│   │   └── helpers.py        ← Shared utilities\n"
"│   ├── services/\n"
"│   │   └── starter_wardrobe.py ← Auto-load default clothes\n"
"│   ├── utils/\n"
"│   │   ├── weather_api.py    ← OpenWeatherMap API client\n"
"│   │   └── image_utils.py    ← File save/validate utilities\n"
"│   └── data/\n"
"│       ├── preloaded_wardrobe.json  ← 32 starter garments\n"
"│       ├── trends.json              ← Seasonal fashion trends\n"
"│       ├── color_mapping.json       ← 100+ colour name→hex\n"
"│       └── skin_tone_palettes.json  ← Colour palettes per skin\n"
"├── frontend/\n"
"│   ├── app.py               ← Home/landing page\n"
"│   └── pages/               ← 7 feature pages\n"
"├── adorkable.db             ← SQLite database file\n"
"├── requirements.txt         ← Python libraries needed\n"
"└── .env                     ← Secret keys (not in Git)\n"
    )

    pb(doc)
    h2(doc,"A.3  Step-by-Step: How a Recommendation is Generated")
    body(doc,"When you click 'Get My Outfit', here is EXACTLY what happens inside the computer:")
    steps=[
        ("Step 1 — Button click","Streamlit frontend sends HTTP POST to http://localhost:8006/api/v1/recommend with JSON body: {occasion: 'Casual', style_pref: 'Western'}"),
        ("Step 2 — Identity check","FastAPI extracts the Bearer JWT token from the request header. python-jose verifies the signature and expiry. The user ID is extracted."),
        ("Step 3 — Load wardrobe","SQLAlchemy queries garment_items table for all garments owned by this user. If fewer than 3 items exist, starter catalog items are added temporarily."),
        ("Step 4 — Fetch weather","weather_api.py calls OpenWeatherMap API with the user's city. Returns temp_c (e.g. 28.0), condition ('Sunny'), humidity (65). Falls back to 22°C/Clear if API unavailable."),
        ("Step 5 — Load profile","SQLAlchemy loads user_profile: skin_tone='Medium', skin_undertone='Warm', body_shape='Pear'."),
        ("Step 6 — Generate candidates","outfit_constraints.py creates ALL valid combinations: every (top+bottom) pair, every dress alone. Eastern style → only traditional_top+traditional_bottom pairs."),
        ("Step 7 — Score candidates","outfit_scorer.py scores every candidate out of 105: colour harmony (30pts) + skin flattery (20pts) + body shape (20pts) + weather (20pts) + occasion (10pts) + trend bonus (5pts)."),
        ("Step 8 — Stochastic selection","stochastic_selector.py takes the top 5 candidates, applies softmax to compute probabilities, then randomly picks one (best outfit chosen ~95% of the time)."),
        ("Step 9 — Hijab addition","helpers.py: if user is female, a hijab is randomly selected from the wardrobe hijab items and added to the outfit."),
        ("Step 10 — Log to database","OutfitLog record saved with today's date, score, occasion, and garment IDs."),
        ("Step 11 — Send JSON response","FastAPI returns: {score: 87.4, top: {...}, bottom: {...}, why_this_suits_you: '...', weather_explanation: '...'}"),
        ("Step 12 — Display on screen","Streamlit renders garment images in cards, shows the score bar, the explanation text, and weather context."),
    ]
    for title,desc in steps:
        h4(doc,title)
        body(doc,desc)

    pb(doc)
    h2(doc,"A.4  How the AI Detects Your Skin Tone")
    body(doc,"When you upload a selfie, this chain of events happens automatically:")
    substeps=[
        ("Load photo","OpenCV reads the image file into a NumPy array — a 3D grid of numbers (height × width × 3 colour channels). Each pixel has a Red, Green, Blue value from 0 to 255."),
        ("Face detection","The BlazeFace TFLite model (pre-trained by Google) scans the image and outputs a bounding box: (x, y, width, height) of the face in normalised 0–1 coordinates."),
        ("Convert to pixel coords","x_pixel = x_normalised × image_width. y_pixel = y_normalised × image_height."),
        ("Crop cheek region","The code cuts out a small rectangular region from your left/right cheeks — NOT the nose (shadow) or forehead (highlight), as those distort the colour."),
        ("Average RGB","NumPy computes the mean Red, Green, Blue values across all pixels in the cheek crop. Example result: R=180, G=140, B=110."),
        ("Classify skin tone level","brightness = (R+G+B)/3. Then:\n   brightness > 220 → 'Very Fair'\n   185–220 → 'Fair'\n   155–185 → 'Light Medium'\n   125–155 → 'Medium'\n   90–125 → 'Medium Dark'\n   < 90 → 'Dark'"),
        ("Classify undertone","R/G ratio:\n   > 1.15 → 'Warm' (golden, peachy tones)\n   < 0.95 → 'Cool' (pink, bluish tones)\n   between → 'Neutral'"),
        ("Save to database","skin_tone and skin_undertone saved to user_profiles table via SQLAlchemy."),
    ]
    for title,desc in substeps:
        h4(doc,title); body(doc,desc)

    pb(doc)
    h2(doc,"A.5  How the AI Detects Your Body Shape")
    body(doc,"When you upload a full-body photo, the PoseLandmarker TFLite model finds 33 body joint points (a digital skeleton).")
    substeps2=[
        ("Pose detection","MediaPipe PoseLandmarker outputs 33 landmarks with normalised (x,y,z) coordinates. Think of it as placing dots on all your joints."),
        ("Shoulder width","Distance between point 11 (left shoulder) and point 12 (right shoulder), converted to pixels."),
        ("Hip width","Distance between point 23 (left hip) and point 24 (right hip), converted to pixels."),
        ("Ratio calculation","ratio = shoulder_width ÷ hip_width. This single number characterises your body proportions."),
        ("Shape classification",
         "ratio > 1.15  → Inverted Triangle (broad shoulders)\n"
         "ratio < 0.85  → Pear (wider hips)\n"
         "ratio ≈ 1.0 + waist visibly narrower → Hourglass\n"
         "ratio ≈ 1.0 + uniform width → Rectangle\n"
         "waist wider than hips/shoulders → Apple\n"
         "strong upper body, moderate taper → Athletic"),
        ("Save to database","body_shape saved to user_profiles table."),
    ]
    for title,desc in substeps2:
        h4(doc,title); body(doc,desc)

    table(doc,
        ["Body Shape","Best Silhouettes","Avoid"],
        [
            ["Hourglass","Fitted, belted, wrap, A-line","Boxy, oversized, shapeless"],
            ["Pear","A-line, wide-leg, structured shoulder","Pencil skirt, skinny bottom"],
            ["Inverted Triangle","Flared bottom, wide-leg, V-neck","Shoulder pads, puffed sleeves"],
            ["Rectangle","Belted, layered, peplum, ruched","Straight cut, boxy"],
            ["Apple","Empire waist, wrap, V-neck, flowing tunic","Clingy knits at midsection"],
            ["Athletic","Balanced separates, slim-leg, column dress","Overly padded shoulders"],
        ]
    )

    pb(doc)
    h2(doc,"A.6  How Outfit Scoring Works — All 5 Dimensions")
    info(doc,"Think of 5 judges each giving marks. Their combined score out of 105 determines which outfit wins.","E8EAF6")

    h3(doc,"Judge 1 — Colour Harmony (max 30 points)")
    body(doc,"Colours can be placed on a circular wheel (0°–360°). The angular distance between two colours determines their harmony type:")
    table(doc,["Harmony Type","Angle","Points Awarded","Real Example"],[
        ["Complementary","~180° apart","30/30","Navy Blue + Orange"],
        ["Triadic","~120° apart","27/30","Red + Blue + Yellow"],
        ["Neutral pair","Either is Black/White/Beige/Grey","24/30","Any colour + Black"],
        ["Analogous","Within 30°","21/30","Blue + Blue-Green"],
        ["Monochromatic","Same hue, different shade","18/30","Navy + Sky Blue"],
        ["No harmony","None of the above","5/30","Random mismatch"],
    ])

    h3(doc,"Judge 2 — Skin Tone Flattery (max 20 points)")
    table(doc,["Your Undertone","Flattering Colours","Avoid"],[
        ["Warm","Coral, gold, olive, peach, rust, camel, warm brown","Icy pastels, stark silver-grey"],
        ["Cool","Navy, emerald, rose, lilac, icy pink, burgundy, silver","Orange, yellow-green, warm brown"],
        ["Neutral","Most colours; true primaries especially","Extreme neons"],
    ])

    h3(doc,"Judge 3 — Body Shape Suitability (max 20 points)")
    body(doc,"Each garment has a style label. The scorer checks if that style is in your body shape's 'best' list (→ 1.0 × 20 = 20 pts) or 'avoid' list (→ 0.2 × 20 = 4 pts).")

    h3(doc,"Judge 4 — Weather Suitability (max 20 points)")
    table(doc,["Temperature","Required Fabric","Score if Perfect Match"],[
        ["Above 28°C (Hot)","light (silk, chiffon, linen)","20/20"],
        ["22–28°C (Warm)","light-medium (cotton, rayon)","20/20"],
        ["15–22°C (Mild)","medium (cotton, light denim)","20/20"],
        ["10–15°C (Cool)","medium-heavy (light wool, corduroy)","20/20"],
        ["Below 10°C (Cold)","heavy (wool, tweed, fleece)","20/20"],
    ])
    body(doc,"Off by one weight level → 12/20. Off by two levels → 4/20. Rain + no outerwear → -6 penalty.")

    h3(doc,"Judge 5 — Occasion Match (max 10 points)")
    body(doc,"Each garment has occasion tags like ['Casual','Work','Party']. Full match = 10/10. Partial match = 5/10. No match = 0/10.")

    h3(doc,"Bonus — Trending Badge (+5 points)")
    body(doc,"trends.json lists current seasonal fashion colours and styles. If your outfit matches → +5 points + 🔥 badge shown on screen.")

    pb(doc)
    h2(doc,"A.7  How Variety is Guaranteed — The Softmax Algorithm")
    body(doc,"If the AI always picked the #1 outfit, you would see the same clothes every day. Softmax Weighted Random Selection prevents this:")
    body(doc,"Example — Top 3 candidate scores: 85, 82, 78")
    code(doc,
"Step 1: Subtract max (85) for numerical stability:\n"
"        85-85=0,  82-85=-3,  78-85=-7\n\n"
"Step 2: Apply e^x:\n"
"        e^0=1.000,  e^-3=0.050,  e^-7=0.001\n\n"
"Step 3: Divide by sum (1.051) → probabilities:\n"
"        Outfit 1: 95.1%  (best, chosen most often)\n"
"        Outfit 2:  4.7%  (good, chosen sometimes)\n"
"        Outfit 3:  0.1%  (acceptable, rarely chosen)\n\n"
"Step 4: Python picks randomly using these probabilities.\n"
"        → Quality maintained, repetition avoided."
    )

    pb(doc)
    h2(doc,"A.8  How the Weekly Planner Works")
    h3(doc,"2-Day Cooldown Rule")
    body(doc,"A garment worn on Day 0 (Monday) cannot appear on Day 1 (Tuesday). It is allowed again from Day 2 (Wednesday). This prevents wearing the same shirt two days in a row.")
    code(doc,
"used_items_with_day = []   # tracks (garment_id, day_number)\n\n"
"For day_idx = 3 (Wednesday):\n"
"    blocked = [id for (id, day) in used_items_with_day\n"
"               if 3 - day < 2]  # days 1,2 are blocked\n"
"    select outfit NOT using any blocked garments"
    )
    h3(doc,"Bottom Variety Boost")
    body(doc,"Outfits using a bottom not worn recently get +15 bonus points before selection. This strongly encourages picking different trousers/skirts each day.")
    h3(doc,"Default Occasions Pattern")
    code(doc,
"Monday    → Casual\nTuesday   → Academic\nWednesday → Casual\n"
"Thursday  → Formal\nFriday    → Casual\nSaturday  → Casual\nSunday    → Formal"
    )

    pb(doc)
    h2(doc,"A.9  How Eastern Style Rules are Enforced")
    table(doc,["Category","Style","Example Garments"],[
        ["top","Western","T-shirt, blouse, shirt, sweater"],
        ["bottom","Western","Jeans, trousers, skirt"],
        ["traditional_top","Eastern","Kurta, kameez"],
        ["traditional_bottom","Eastern","Shalwar, salwar, churidar"],
        ["dress","Any","Western dress, abaya, sari"],
        ["hijab","Islamic","All hijab styles"],
    ])
    body(doc,"The enforcement rule (a single if-statement in the outfit generation loop):")
    code(doc,
"for top in all_tops:\n"
"  for bottom in all_bottoms:\n"
"    if style_pref == 'Eastern':\n"
"      if not ('traditional' in top.category\n"
"              and 'traditional' in bottom.category):\n"
"          continue  # skip — mixed Western+Eastern pair\n"
"    score_this_combination(top, bottom)"
    )
    body(doc,"Result: Eastern style users ONLY see kurta+shalwar, kameez+churidar etc. Western t-shirt with shalwar is NEVER suggested.")

    pb(doc)
    h2(doc,"A.10  How the Database Stores Everything")
    table(doc,["Table","What It Stores","Key Columns"],[
        ["users","User accounts","id, email, hashed_password, gender, city"],
        ["user_profiles","AI analysis results","skin_tone, skin_undertone, body_shape"],
        ["garment_items","All uploaded clothes","category, color_hex, fabric_weight, occasion_tags, wear_count"],
        ["outfit_logs","Recommendation history","worn_date, score, top_id, bottom_id, dress_id, occasion"],
    ])
    bullet(doc,"Passwords are NEVER stored as plain text — only bcrypt hashes")
    bullet(doc,"One user → many garments (one-to-many relationship)")
    bullet(doc,"One user → one profile (one-to-one relationship)")
    bullet(doc,"One user → many outfit logs (one-to-many relationship)")
    pb(doc)


def build_part_b(doc):
    h1(doc,"PART B — TECHNOLOGY STACK: EVERY LIBRARY EXPLAINED",NAVY)

    h2(doc,"B.1  Full Library Reference")

    libs=[
        ("FastAPI","0.115","Backend Web Framework",
         "Creates the server that listens for requests from the frontend. Uses Python's async/await for handling multiple users simultaneously. Automatically creates a testing UI at /docs.",
         "Fastest Python web framework; built-in data validation via Pydantic; native async support."),
        ("Uvicorn","0.32","ASGI Web Server",
         "The actual program that runs FastAPI. Listens on port 8006 and passes HTTP requests to FastAPI. Like the engine of a car — FastAPI is the steering wheel.",
         "Required by FastAPI; uses uvloop for ultra-fast async I/O."),
        ("Streamlit","1.40","Frontend Framework",
         "Converts Python scripts into interactive web pages. You write Python code like st.button() and Streamlit renders a real button in the browser. No HTML or JavaScript needed.",
         "Allows the team to build the UI in Python only; built-in reactive state management."),
        ("SQLAlchemy 2.0","2.0.36","Database ORM",
         "Lets you work with the database using Python classes instead of SQL queries. Define a class → SQLAlchemy creates the table. Add an object → SQLAlchemy saves the row.",
         "Version 2.0 adds full async support needed for FastAPI."),
        ("aiosqlite","0.20","Async SQLite Driver",
         "Makes SQLite work with async/await. Without it, every database operation would block the server and slow down all other users.",
         "Required for async SQLAlchemy with SQLite."),
        ("MediaPipe","0.10.18","Computer Vision AI",
         "Google's library for running pre-trained AI vision models. Used for BlazeFace (face detection) and PoseLandmarker (body skeleton). Models are downloaded automatically.",
         "State-of-the-art accuracy; runs on CPU; no GPU required; free and open source."),
        ("OpenCV (cv2)","4.10","Image Processing",
         "Reads image files into NumPy arrays, converts colour formats (BGR↔RGB), resizes images, and crops regions. Every image in the pipeline passes through OpenCV.",
         "Universal standard; works with NumPy natively; used by MediaPipe."),
        ("scikit-learn","1.5.2","Machine Learning",
         "Provides KMeans clustering for colour extraction. Groups all pixels in a garment image into 5 colour clusters; the largest cluster is the dominant colour.",
         "Reliable, well-documented; perfect NumPy integration."),
        ("NumPy","2.0.2","Numerical Computing",
         "Images are NumPy arrays (height×width×3). All pixel averaging, distance calculation, and array manipulation uses NumPy operations.",
         "Foundation of all scientific Python; required by every other library here."),
        ("Pillow","11.0","Image File I/O",
         "Reads and validates image files (JPEG, PNG, WebP). Validates uploads before processing.",
         "Best format support; integrates with Streamlit for display."),
        ("TensorFlow","2.18","AI Model Runtime",
         "Runs in the background as the computational engine for MediaPipe's TFLite models. We do not write TensorFlow code — it executes automatically when MediaPipe loads a model.",
         "Required by MediaPipe."),
        ("python-jose[cryptography]","3.3","JWT Authentication",
         "Creates and verifies JWT tokens (digital ID cards). Encodes user ID + expiry time into a signed string that cannot be forged.",
         "RFC-compliant JWT; HS256 signing; industry standard for API authentication."),
        ("passlib[bcrypt]","1.7.4","Password Security",
         "Hashes passwords using bcrypt — a one-way mathematical scramble with random salt. The original password can never be recovered from the hash.",
         "OWASP-recommended; adaptive cost factor prevents brute-force attacks."),
        ("httpx","0.27","HTTP Client",
         "Used by the Streamlit frontend to send HTTP requests to the FastAPI backend. When you click a button, httpx sends the POST request and receives the JSON response.",
         "Better timeout control than requests library; crucial for AI endpoints that take seconds."),
        ("plotly","5.24","Interactive Charts",
         "Renders the analytics dashboard charts: colour distribution donut, wear frequency bar chart, outfit score trend line. Charts are interactive — hover for values.",
         "Streamlit-native integration via st.plotly_chart."),
        ("python-dotenv","1.0.1","Config Management",
         "Reads the .env file and loads API keys and settings as environment variables so secrets are never hardcoded in source code.",
         "Security best practice; prevents accidental secret exposure on GitHub."),
        ("requests","2.32","HTTP Client (Backend)",
         "Used in weather_api.py to call the OpenWeatherMap REST API. Simple synchronous GET request with city name and API key.",
         "Simple; appropriate for a single synchronous background task."),
    ]
    for name,ver,role,what,why in libs:
        h3(doc,f"{name}  (v{ver})")
        lv(doc,"Role",role)
        lv(doc,"What it does",what)
        lv(doc,"Why chosen",why)
        doc.add_paragraph()

    pb(doc)
    h2(doc,"B.2  How Libraries Work Together — Complete Data Journeys")
    info(doc,"Three complete journeys from user action to final result.","E8F5E9")

    journeys=[
        ("JOURNEY 1: User uploads a garment photo",[
            "Pillow validates the file is a real image (not a text file pretending to be JPEG)",
            "image_utils.py saves it to disk as uploads/<uuid>.png using a random UUID filename",
            "OpenCV reads the file as BGR NumPy array → converts to RGB",
            "Resize to 150×150 pixels for speed",
            "NumPy reshapes to (22500, 3) pixel array",
            "scikit-learn KMeans(n_clusters=5) finds 5 colour groups",
            "NumPy counts pixels per group; largest group = dominant colour",
            "rgb→hex conversion; color_mapping.json lookup → fashion colour name",
            "SQLAlchemy saves new GarmentItem row to adorkable.db",
        ]),
        ("JOURNEY 2: User uploads a selfie",[
            "Pillow validates and saves the file",
            "OpenCV reads as BGR NumPy array → converts to RGB",
            "MediaPipe (TensorFlow TFLite backend) runs BlazeFace → face bounding box",
            "NumPy crops cheek region using bounding box coordinates",
            "NumPy .mean(axis=(0,1)) averages all cheek pixels → single RGB value",
            "Custom Python code: brightness → 6-level skin tone",
            "Custom Python code: R/G ratio → Warm/Cool/Neutral undertone",
            "SQLAlchemy saves to user_profiles table",
        ]),
        ("JOURNEY 3: User requests a daily outfit",[
            "httpx POST from Streamlit → FastAPI endpoint",
            "python-jose decodes and verifies JWT Bearer token",
            "SQLAlchemy queries garment_items (user's wardrobe)",
            "requests calls OpenWeatherMap API → weather dict",
            "SQLAlchemy queries user_profiles → skin tone + body shape",
            "outfit_constraints.py generates all valid top+bottom and dress combinations",
            "outfit_scorer.py: color_theory.py + weather_rules.py + config.py rules → score each",
            "stochastic_selector.py: softmax over top-5 → random weighted pick",
            "helpers.py: if female → add hijab from wardrobe",
            "SQLAlchemy saves OutfitLog",
            "FastAPI returns JSON response",
            "Streamlit renders garment image cards + score + explanation",
        ]),
    ]
    for title,steps in journeys:
        h3(doc,title)
        for i,s in enumerate(steps,1):
            bullet(doc,f"Step {i}: {s}")
        doc.add_paragraph()
    pb(doc)


if __name__=="__main__":
    doc=Document()
    build_title(doc)
    build_toc(doc)
    build_part_a(doc)
    build_part_b(doc)
    out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"_master_p1.docx")
    doc.save(out)
    print(f"Part 1 saved → {out}")
