"""
Adorkable AI — MASTER DOCUMENT BUILDER
Assembles all 3 parts into one final DOCX file.

Run:  python docs/BUILD_MASTER_DOC.py
Output: docs/Master_Project_Guide.docx
"""
import os, sys

# Add project root to path
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Import all three parts
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from build_master_p1 import build_title, build_toc, build_part_a, build_part_b
from build_master_p2 import build_part_c
from build_master_p3 import build_part_d, build_part_e, build_part_f

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Master_Project_Guide.docx")


def set_page_margins(doc):
    """Set comfortable page margins."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)


def add_header_footer(doc):
    """Add page header and footer."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for section in doc.sections:
        # Header
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        run = p.add_run("Adorkable AI — Master Project Guide  |  AI-Powered Fashion Stylist")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        run2 = fp.add_run("Confidential — Academic Project Documentation  |  Department of CS & AI  |  2025-2026")
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run2.italic = True
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    print("Building Adorkable AI Master Project Guide...")
    print("=" * 55)

    doc = Document()
    set_page_margins(doc)
    add_header_footer(doc)

    print("[1/8] Building Title Page...")
    build_title(doc)

    print("[2/8] Building Table of Contents...")
    build_toc(doc)

    print("[3/8] Building Part A — How the System Works...")
    build_part_a(doc)

    print("[4/8] Building Part B — Technology Stack...")
    build_part_b(doc)

    print("[5/8] Building Part C — Every File Explained...")
    build_part_c(doc)

    print("[6/8] Building Part D — Complete Data Flows...")
    build_part_d(doc)

    print("[7/8] Building Part E — Algorithms Explained...")
    build_part_e(doc)

    print("[8/8] Building Part F — How to Explain to Anyone...")
    build_part_f(doc)

    print(f"\nSaving document to:\n  {OUT}")
    doc.save(OUT)

    size_kb = os.path.getsize(OUT) // 1024
    print(f"\n✅ SUCCESS! Master_Project_Guide.docx saved ({size_kb} KB)")
    print("\nDocument contains:")
    print("  • Title Page")
    print("  • Table of Contents (searchable headings)")
    print("  • Part A: How the Whole System Works (10 sections)")
    print("  • Part B: All 16 Libraries Explained + Data Journey Maps")
    print("  • Part C: Every File Explained — 34 files, all functions")
    print("  • Part D: 4 Complete Step-by-Step Data Flows")
    print("  • Part E: 7 Algorithms in Plain English")
    print("  • Part F: Q&A Prep + How to Explain to Anyone")
    print("\nTo open: double-click docs/Master_Project_Guide.docx")
    print("In Word, use Ctrl+F to search any section title.")


if __name__ == "__main__":
    main()
