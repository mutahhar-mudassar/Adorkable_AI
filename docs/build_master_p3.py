"""Part 3: Parts D, E, F — Data Flow, Algorithms, How to Explain"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1A,0x1A,0x5E); BLUE=RGBColor(0x00,0x5F,0xA3); TEAL=RGBColor(0x00,0x7A,0x7A)
DARK=RGBColor(0x22,0x22,0x22); GREEN=RGBColor(0x19,0x6F,0x3D); ORANGE=RGBColor(0xCA,0x6F,0x1E)

def sp(p,f="F4F4F4"):
    pp=p._p.get_or_add_pPr(); s=OxmlElement("w:shd")
    s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),f); pp.append(s)

def h1(d,t):
    p=d.add_heading(t,1); r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True; p.paragraph_format.space_before=Pt(16)

def h2(d,t):
    p=d.add_heading(t,2); r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=BLUE; r.font.size=Pt(14); r.bold=True; p.paragraph_format.space_before=Pt(12)

def h3(d,t):
    p=d.add_heading(t,3); r=p.runs[0] if p.runs else p.add_run(t)
    r.font.color.rgb=TEAL; r.font.size=Pt(12); r.bold=True

def h4(d,t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=GREEN
    p.paragraph_format.space_before=Pt(8)

def body(d,t):
    p=d.add_paragraph(t)
    for r in p.runs: r.font.size=Pt(11); r.font.color.rgb=DARK
    p.paragraph_format.space_after=Pt(4)

def bul(d,t,lv=0):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t)
    r.font.size=Pt(11); r.font.color.rgb=DARK; p.paragraph_format.left_indent=Inches(0.3+lv*0.3)

def num(d,t):
    p=d.add_paragraph(style="List Number"); r=p.add_run(t)
    r.font.size=Pt(11); r.font.color.rgb=DARK

def code(d,t):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
    r=p.add_run(t); r.font.name="Courier New"; r.font.size=Pt(9); r.font.color.rgb=DARK; sp(p)

def info(d,t,f="FFF8E1"):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3)
    r=p.add_run(t); r.font.size=Pt(10); r.italic=True; r.font.color.rgb=RGBColor(0x5D,0x40,0x00); sp(p,f)

def tbl(d,headers,rows):
    t=d.add_table(rows=1+len(rows),cols=len(headers)); t.style="Table Grid"
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; r=c.paragraphs[0].runs[0]
        r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        tc=c._tc; tcp=tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
        s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),"1A1A5E"); tcp.append(s)
    for ri,rd in enumerate(rows):
        row=t.rows[ri+1].cells
        for ci,ct in enumerate(rd): row[ci].text=str(ct)
        fill="F0F4FF" if ri%2==0 else "FFFFFF"
        for ci in range(len(headers)):
            tc=t.rows[ri+1].cells[ci]._tc; tcp=tc.get_or_add_tcPr(); s=OxmlElement("w:shd")
            s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),fill); tcp.append(s)
    d.add_paragraph()

def pb(d): d.add_page_break()


def build_part_d(doc):
    h1(doc,"PART D — COMPLETE DATA FLOW WALKTHROUGH")
    info(doc,"Follow data from the moment a user opens the browser to the final outfit displayed on screen.","E3F2FD")

    h2(doc,"D.1  Complete User Registration Flow")
    body(doc,"What happens when a brand new user signs up:")
    steps=[
        "User opens http://localhost:8501 → Streamlit serves app.py (the landing page)",
        "User clicks 'Register' → navigates to 1_Register_Login.py",
        "User fills in: email='sara@example.com', password='mypass', gender='female', city='Lahore'",
        "User clicks 'Create Account'",
        "Streamlit: httpx.post('http://localhost:8006/api/v1/auth/register', json={...})",
        "FastAPI receives the request → auth.py register endpoint runs",
        "Check: SELECT * FROM users WHERE email='sara@example.com' → returns nothing (email free)",
        "passlib.bcrypt: hash_password('mypass') → '$2b$12$...hash...'",
        "SQLAlchemy: INSERT INTO users (email, hashed_password, gender, city) VALUES (...)",
        "python-jose: create_access_token({'sub': '1'}) → JWT string",
        "FastAPI returns: {'access_token': 'eyJ...', 'token_type': 'bearer'}",
        "Streamlit: st.session_state.token = 'eyJ...' (stored for all future requests)",
        "Streamlit: st.switch_page('pages/2_My_Profile.py')",
    ]
    for i,s in enumerate(steps,1): num(doc,f"Step {i}: {s}")
    pb(doc)

    h2(doc,"D.2  Complete Garment Upload Flow")
    body(doc,"What happens when a user uploads a photo of their shirt:")
    steps=[
        "User is on the Wardrobe page (3_My_Wardrobe.py), Upload tab",
        "User selects: Category='top', Style='casual', Occasions=['Casual','Work'], Fabric='medium'",
        "User clicks 'Choose file' and selects shirt.jpg from their computer",
        "User clicks 'Upload Garment'",
        "Streamlit: httpx.post('/api/v1/wardrobe/upload', files={'file': shirt_bytes}, data={metadata})",
        "FastAPI: wardrobe.py upload endpoint receives the request",
        "is_valid_image(file_bytes) → Pillow opens successfully → valid",
        "save_uploaded_image(bytes, user_id=1) → saves as 'uploads/f3a9b2c1.jpg' → returns path",
        "extract_dominant_color('uploads/f3a9b2c1.jpg'):",
        "  → OpenCV reads image as BGR NumPy array (e.g. 400×300×3)",
        "  → Convert BGR→RGB",
        "  → Resize to 150×150 (= 22,500 pixels)",
        "  → Reshape to (22500, 3) array",
        "  → KMeans(n_clusters=5).fit(pixels) — finds 5 colour clusters",
        "  → Largest cluster centroid: RGB=(210, 220, 255) → hex='#D2DCFF'",
        "  → hex_to_color_name('#D2DCFF') → Euclidean distance search → 'Lavender Blue'",
        "classify_garment('top','casual',['Casual','Work'],'medium') → attribute dict",
        "SQLAlchemy: INSERT INTO garment_items (user_id=1, image_path='uploads/f3.jpg', category='top', dominant_color='Lavender Blue', ...) VALUES (...)",
        "FastAPI returns: {'id':5, 'category':'top', 'dominant_color':'Lavender Blue', 'color_hex':'#D2DCFF', ...}",
        "Streamlit: shows success message + garment thumbnail",
    ]
    for i,s in enumerate(steps,1): num(doc,f"Step {i}: {s}")
    pb(doc)

    h2(doc,"D.3  Complete Selfie Analysis Flow")
    body(doc,"What happens when a user uploads their selfie photo:")
    steps=[
        "User is on Profile page (2_My_Profile.py), Selfie section",
        "User uploads selfie.jpg",
        "Streamlit: httpx.post('/api/v1/profile/selfie', files={'file': selfie_bytes})",
        "FastAPI: profile.py selfie endpoint receives request",
        "save_uploaded_image(bytes, user_id=1, filename='selfie.jpg') → 'uploads/abc123_selfie.jpg'",
        "analyse_skin_tone('uploads/abc123_selfie.jpg') calls:",
        "  → initialise_face_detector() (loads BlazeFace TFLite model)",
        "  → OpenCV: img = cv2.imread(path); img_rgb = cv2.cvtColor(img, BGR2RGB)",
        "  → MediaPipe (TensorFlow TFLite): face_detector.detect(img_rgb)",
        "  → Detection found! bbox = (x=0.32, y=0.18, w=0.25, h=0.35) [normalised]",
        "  → Convert to pixels: x_px=102, y_px=58, w_px=80, h_px=112",
        "  → Crop left cheek: img_array[85:120, 108:135] (specific pixel range)",
        "  → NumPy: cheek_avg = cheek_pixels.mean(axis=(0,1)) → [R=175, G=142, B=118]",
        "  → brightness = (175+142+118)/3 = 145 → 'Medium' skin tone",
        "  → R/G ratio = 175/142 = 1.23 > 1.15 → 'Warm' undertone",
        "  → returns {'skin_tone':'Medium', 'skin_undertone':'Warm', 'confidence':0.91}",
        "SQLAlchemy: UPDATE user_profiles SET skin_tone='Medium', skin_undertone='Warm' WHERE user_id=1",
        "FastAPI returns: {'skin_tone':'Medium', 'skin_undertone':'Warm', 'flattering_colors':['Coral','Gold','Olive',...]}",
        "Streamlit: displays skin tone badge + flattering colour swatches",
    ]
    for i,s in enumerate(steps,1): num(doc,f"Step {i}: {s}")
    pb(doc)

    h2(doc,"D.4  Complete Daily Outfit Recommendation Flow")
    body(doc,"The full detailed flow for generating a daily outfit recommendation:")
    steps=[
        "User is on Daily Outfit page (4_Daily_Outfit.py)",
        "User selects: Occasion='Casual', Style='Western'",
        "User clicks 'Get My Outfit'",
        "Streamlit: httpx.post('/api/v1/recommend', json={'occasion':'Casual','style_pref':'Western','reimagine_step':0}, headers={'Authorization':'Bearer eyJ...'})",
        "FastAPI: recommendations.py → get_current_user() dependency runs",
        "  → Extract Bearer token from Authorization header",
        "  → python-jose: decode_token(token) → payload={'sub':'1','exp':1748...}",
        "  → Token not expired ✓",
        "  → SQLAlchemy: SELECT * FROM users WHERE id=1 → User object returned",
        "get_user_wardrobe(db, user_id=1) → returns 12 GarmentItem objects",
        "Wardrobe check: 12 items ≥ 3 → no need for starter items",
        "get_current_weather('Lahore'):",
        "  → requests.get('https://api.openweathermap.org/data/2.5/weather?q=Lahore&...')",
        "  → Response: {'main':{'temp':32.4,'humidity':68},'weather':[{'main':'Clear'}]}",
        "  → returns {'temp_c':32.4, 'condition':'Clear', 'humidity':68}",
        "SQLAlchemy: SELECT * FROM user_profiles WHERE user_id=1 → UserProfile object",
        "  → skin_tone='Medium', skin_undertone='Warm', body_shape='Pear'",
        "generate_wardrobe_outfit_candidates(wardrobe=12items, profile, weather, 'Casual', 'Western'):",
        "  → tops=[3 items], bottoms=[3 items], dresses=[2 items]",
        "  → Top+Bottom combos: 3×3=9 pairs generated and scored",
        "  → Dress combos: 2 dresses scored individually",
        "  → Total 11 candidates scored",
        "All 11 candidates scored:",
        "  → Best: white_tshirt+blue_jeans = 73.4pts",
        "  → 2nd: coral_blouse+white_trousers = 69.2pts",
        "  → etc.",
        "weighted_random_select(candidates, top_n=5, skip=0):",
        "  → Takes top 5 by score",
        "  → softmax([73.4, 69.2, 65.1, 61.8, 58.3]) → probs=[0.71, 0.17, 0.07, 0.03, 0.01]",
        "  → random.choices picks... white_tshirt+blue_jeans (71% chance this time)",
        "select_hijab_for_female('female', items) → User is female → finds hijab_black in wardrobe",
        "SQLAlchemy: INSERT INTO outfit_logs (user_id=1, top_id=3, bottom_id=7, score=73.4, occasion='Casual', worn_date=today)",
        "UPDATE garment_items SET wear_count=wear_count+1 WHERE id IN (3,7)",
        "Build response dict: {score:73.4, top:{id:3,category:'top',color:'White',...}, bottom:{id:7,...}, hijab:{...}, why_this_suits_you:'White harmonises beautifully with Navy in a classic Complementary palette...', weather_explanation:'At 32.4°C it is hot — the lightweight cotton fabric is ideal', trending:False}",
        "FastAPI returns JSON response",
        "Streamlit: renders 3 garment image cards side by side, score progress bar, explanation text",
    ]
    for i,s in enumerate(steps,1): num(doc,f"Step {i}: {s}")
    pb(doc)


def build_part_e(doc):
    h1(doc,"PART E — ALGORITHMS EXPLAINED IN PLAIN ENGLISH")
    info(doc,"Detailed explanation of every algorithm used — no mathematics background required.","E8F5E9")

    h2(doc,"E.1  KMeans Clustering — How Garment Colours Are Found")
    body(doc,"Imagine you have a bag of 22,500 marbles, each painted a slightly different colour. KMeans organises them into 5 groups by colour similarity.")
    h3(doc,"Step-by-step:")
    num(doc,"Step 1 — Pick 5 random starting points (called 'centroids') from the marble pile")
    num(doc,"Step 2 — For every marble, measure its distance (in RGB space) to all 5 centroids. Assign the marble to the nearest centroid's group.")
    num(doc,"Step 3 — Move each centroid to the mathematical centre (average) of its assigned marbles")
    num(doc,"Step 4 — Repeat Steps 2–3 until no marble changes groups (convergence)")
    num(doc,"Step 5 — Count how many marbles are in each group. The biggest group = dominant colour. Its centroid = the dominant colour's RGB value.")
    body(doc,"Why 5 clusters? More clusters give more detail but slower computation. 5 is enough to identify the main colour, secondary colours, and background/white.")
    info(doc,"Real example: A navy blue shirt photo has mostly dark blue pixels, some lighter blue, some white collar, some shadow. KMeans group 1 (largest) = navy blue pixels → dominant colour = Navy Blue.","E3F2FD")

    h2(doc,"E.2  MediaPipe BlazeFace — How Faces Are Detected")
    body(doc,"BlazeFace is a neural network — a mathematical function with millions of learned parameters (weights). It was trained by showing it millions of photos with manually labelled face locations.")
    h3(doc,"How it works on a new photo:")
    num(doc,"The photo is divided into a grid of overlapping windows at multiple scales (different zoom levels)")
    num(doc,"Each window is passed through the neural network (hundreds of mathematical multiplications and additions)")
    num(doc,"The network outputs: probability that this window contains a face (0.0–1.0) + bounding box offset adjustments")
    num(doc,"Windows with probability > 0.5 are considered face detections")
    num(doc,"Non-Maximum Suppression removes overlapping detections, keeping only the best one")
    num(doc,"Final output: bounding box (x, y, width, height) of the detected face")
    body(doc,"We do NOT train this model. We download the pre-trained .tflite file and use it as a tool. Training took Google weeks on thousands of GPU machines.")

    h2(doc,"E.3  MediaPipe PoseLandmarker — How Body Shape Is Measured")
    body(doc,"PoseLandmarker is another pre-trained neural network that outputs 33 body keypoints.")
    h3(doc,"The 33 keypoints (landmark indices):")
    tbl(doc,["Landmark #","Body Point","Used For"],[
        ["0","Nose","Face direction verification"],
        ["11","Left Shoulder","Shoulder width measurement"],
        ["12","Right Shoulder","Shoulder width measurement"],
        ["13","Left Elbow","Arm length (optional)"],
        ["14","Right Elbow","Arm length (optional)"],
        ["23","Left Hip","Hip width measurement"],
        ["24","Right Hip","Hip width measurement"],
        ["25","Left Knee","Height estimation"],
        ["26","Right Knee","Height estimation"],
    ])
    body(doc,"The shoulder and hip landmarks are the critical ones. The ratio shoulder_width/hip_width is the single most important number for body shape classification.")
    h3(doc,"Why pixel coordinates not centimetres?")
    body(doc,"The image could be any size. Using pixel coordinates and then computing RATIOS (not absolute measurements) means the classification works regardless of whether the photo is taken from 1 metre or 3 metres away. The ratio is scale-invariant.")

    h2(doc,"E.4  Softmax — How Variety Is Mathematically Guaranteed")
    body(doc,"Softmax is a mathematical function that converts any list of numbers into a probability distribution (values that sum to 1.0).")
    h3(doc,"Why not just always pick the best?")
    body(doc,"If we always picked the outfit with the highest score, users would see the same 2–3 outfits repeatedly. Fashion should have variety. But we also can't just pick randomly — that might give a terrible outfit on an important day.")
    h3(doc,"The mathematical formula:")
    code(doc,"P(outfit_i) = e^(score_i) / sum(e^(score_j) for all j in top_n)")
    body(doc,"The exponential function (e^x) amplifies differences: a score of 85 gets e^85 which is astronomically larger than e^78. So even though 78 is close to 85, the probability ratio is extreme.")
    h3(doc,"The numerical stability trick:")
    body(doc,"Direct computation of e^85 causes floating-point overflow in computers. The trick: subtract the maximum score first. e^(85-85)=e^0=1.0, e^(78-85)=e^(-7)=0.001. Same ratios, no overflow.")
    info(doc,"Analogy: Imagine 5 candidates with 'interview scores'. Softmax is like giving the best candidate 71% of all job offers, second-best 17%, etc. The best usually gets the job, but occasionally someone else does — preventing bias.","FFF3E0")

    h2(doc,"E.5  Colour Theory — HSL and the Colour Wheel")
    body(doc,"Computers store colours as RGB (Red, Green, Blue values 0-255). But humans perceive colour relationally — we know 'red and green clash' and 'navy and white look crisp'. To encode this human knowledge, we convert to HSL.")
    h3(doc,"HSL breakdown:")
    tbl(doc,["Component","Range","Meaning","Example"],[
        ["Hue","0°–360°","Colour's position on wheel","0°=Red, 120°=Green, 240°=Blue, 300°=Magenta"],
        ["Saturation","0%–100%","How vivid/grey the colour is","0%=grey, 100%=full vivid colour"],
        ["Lightness","0%–100%","How dark/bright","0%=black, 50%=normal, 100%=white"],
    ])
    h3(doc,"Harmony rules (based on centuries of art and design theory):")
    tbl(doc,["Rule","Hue Angle Difference","Why It Works"],[
        ["Complementary","~180° (opposites)","Maximum contrast, eye-catching"],
        ["Triadic","~120° (triangle)","Balanced, vibrant, energetic"],
        ["Analogous","<30° (neighbours)","Harmonious, soothing, cohesive"],
        ["Monochromatic","<15° (same hue)","Elegant, sophisticated, easy"],
        ["Neutral pairing","One is grey/white/black","Neutrals work with everything"],
    ])
    body(doc,"The colour harmony score directly reflects how 'matched' an outfit looks to the trained eye, encoded mathematically.")

    h2(doc,"E.6  JWT Authentication — How Login Security Works")
    body(doc,"JWT (JSON Web Token) solves a key problem: after logging in, how does the server know who you are on the next request? Traditional solution: store a session in the server's memory. Problem: with many users, this uses lots of memory and doesn't work across multiple servers.")
    h3(doc,"JWT solution — Stateless authentication:")
    num(doc,"At login: server creates a JWT containing your user ID, signed with the SECRET_KEY")
    num(doc,"Server sends JWT to you (the client/browser)")
    num(doc,"Client stores the JWT and sends it with EVERY request in the Authorization header")
    num(doc,"Server verifies the signature — if valid, the user ID inside is trusted")
    num(doc,"No database lookup needed for verification — just cryptographic verification")
    h3(doc,"JWT structure:")
    code(doc,"Header.Payload.Signature\n\nHeader:  {'alg':'HS256','typ':'JWT'} → base64 encoded\nPayload: {'sub':'1','iat':1715000,'exp':1716000} → base64 encoded\nSignature: HMAC_SHA256(header+'.'+payload, SECRET_KEY)")
    body(doc,"The signature is created with HMAC-SHA256 using the SECRET_KEY. Only the server knows this key, so only the server can create valid signatures. A forged token would have an invalid signature and be rejected.")

    h2(doc,"E.7  bcrypt Password Hashing — Why Passwords Are Safe")
    body(doc,"bcrypt is designed to be deliberately slow — it takes about 100ms to hash one password. This is intentional. Fast hashing (like SHA-256) allows attackers to try billions of passwords per second. bcrypt limits them to ~10 attempts per second.")
    h3(doc,"The salting process:")
    code(doc,"salt = generate_random_22_character_string()  # Different every time\nhash = bcrypt(salt + password + salt)  # Salt is embedded in the hash\nstored = '$2b$12$' + salt + hash_bytes")
    body(doc,"The stored string contains: algorithm identifier ($2b$), cost factor ($12$, meaning 2^12 = 4096 iterations), the salt, and the hash — all in one string. At verification time, bcrypt extracts the salt from the stored string and uses it to hash the provided password.")
    info(doc,"Even if two users have the same password, their stored hashes are COMPLETELY different because each has a unique random salt. An attacker who steals the database cannot use pre-computed 'rainbow tables'.","FFE0B2")

    pb(doc)


def build_part_f(doc):
    h1(doc,"PART F — HOW TO EXPLAIN THIS PROJECT TO ANYONE")
    info(doc,"Use these scripts to explain the project confidently to your professors, classmates, or viva committee.","E8F5E9")

    h2(doc,"F.1  The 30-Second Elevator Pitch")
    info(doc,"Use this when someone asks 'What is your project?' in a casual setting.","FFF3E0")
    body(doc,'"Adorkable AI is an AI-powered personal fashion stylist. You upload your wardrobe and a selfie, and the system uses computer vision to detect your skin tone and body shape. Then it recommends outfits from your clothes that are colour-coordinated, weather-appropriate, and suitable for your body type and occasion. It also plans your entire week\'s outfits and makes sure you don\'t repeat clothes too often."')

    h2(doc,"F.2  The 2-Minute Technical Explanation")
    info(doc,"Use this for a professor or technical audience.","FFF3E0")
    body(doc,'"The system has three layers. The frontend is built with Streamlit — a Python web framework that creates the user interface. The backend is built with FastAPI — an async Python web server that handles all API requests. The database is SQLite managed through SQLAlchemy ORM."')
    body(doc,'"For skin tone analysis, we use Google\'s BlazeFace TFLite model to detect the face and crop the cheek region, then classify the average pixel colour into 6 skin tone levels and 3 undertones using brightness and R/G ratio thresholds."')
    body(doc,'"For garment colour extraction, we use KMeans clustering — an unsupervised machine learning algorithm that groups pixels into 5 colour clusters and identifies the dominant one."')
    body(doc,'"The outfit scorer calculates a score out of 105 across 5 dimensions: colour harmony using HSL colour theory, skin tone flattery using curated palettes, body shape suitability using anthropometric rules, weather appropriateness using fabric weight classification, and occasion matching. Variety is ensured using softmax-weighted random selection."')

    h2(doc,"F.3  Questions You Might Be Asked and How to Answer Them")

    qas = [
        ("Q: Why did you use FastAPI instead of Flask or Django?",
         "FastAPI uses Python's async/await system, which allows the server to handle multiple requests simultaneously without blocking. This is crucial because image processing (skin tone, body shape analysis) takes 1-3 seconds. With synchronous frameworks, every user would have to wait in a queue. Additionally, FastAPI automatically generates interactive API documentation at /docs."),
        ("Q: Why SQLite instead of PostgreSQL or MySQL?",
         "SQLite requires zero server configuration — it's just a file on disk. For a demonstration project and single-instance deployment, SQLite is perfectly adequate. The async support via aiosqlite makes it non-blocking. For production scale with multiple users, we would migrate to PostgreSQL using the same SQLAlchemy ORM code — only the DATABASE_URL string would change."),
        ("Q: Did you train the AI models yourself?",
         "We use two pre-trained models from Google's MediaPipe library: BlazeFace for face detection and PoseLandmarker for body pose estimation. These are trained by Google on millions of images and provided as TFLite (TensorFlow Lite) files. We download and use them as tools, similar to how a researcher uses a pre-built calculator. The original model development would require millions of labelled training images and weeks of GPU training time. Our AI contribution is the scoring system, colour theory engine, and recommendation algorithms, which are all custom-built."),
        ("Q: How do you ensure outfit variety?",
         "We use Softmax Weighted Random Selection. Instead of always picking the highest-scored outfit, we convert the top 5 candidate scores into probabilities using the softmax function (e^score / sum(e^scores)). The best outfit gets about a 70% probability, the second about 20%, and so on. We then use Python's random.choices() to make a weighted random selection. This means quality is maintained while variety is guaranteed."),
        ("Q: What is the scoring formula?",
         "Total score = (colour_harmony × 30) + (skin_flattery × 20) + (body_shape × 20) + (weather × 20) + (occasion × 10) + (trending_bonus: 0 or 5). Maximum possible = 105. Each sub-score is normalised to 0.0–1.0 before multiplication. The weights reflect fashion priority: colour harmony is most important (30%), followed equally by personalisation dimensions (skin, body) and contextual fit (weather)."),
        ("Q: How does the Eastern style feature work?",
         "Garments are tagged with categories. Eastern garments have category 'traditional_top' or 'traditional_bottom'. When the user selects Eastern style preference, the outfit generator adds a single if-statement check: 'if style_pref == Eastern, skip any combination where both garments are not tagged as traditional'. This one condition enforces that kurtas only pair with shalwars, never with Western jeans. The hijab is then automatically selected from wardrobe items with category='hijab' for female users."),
        ("Q: How does the weekly planner avoid repetition?",
         "Two mechanisms: 1) Two-day cooldown: a list tracks (garment_id, day_index) pairs. For each day, garments used within the last 2 days are excluded from candidate selection. 2) Bottom variety boost: outfits using bottoms not recently worn receive a +15 point bonus before selection, making the random selector heavily prefer them. The combination ensures different complete outfits across the 7 days."),
        ("Q: What if the user's wardrobe is empty?",
         "The starter_wardrobe.py service checks if the wardrobe is empty before every recommendation. If empty, it loads preloaded_wardrobe.json — a curated catalog of 32 garments — filters by the user's gender, copies the image files to the user's upload directory, and saves GarmentItem rows to the database. The user gets a functional wardrobe immediately, which they can replace with their own uploads over time."),
        ("Q: What is bcrypt and why use it for passwords?",
         "bcrypt is a password hashing algorithm designed to be deliberately slow — about 100ms per hash. This is intentional: attackers trying to crack passwords by testing millions of candidates can only try about 10 per second with bcrypt, versus billions per second with faster algorithms. bcrypt also applies a unique random salt to each password, making rainbow table attacks impossible. Even if two users have identical passwords, their stored hashes are completely different."),
        ("Q: How does the weather integration work?",
         "The backend calls the OpenWeatherMap REST API with the user's city name and API key. The response includes temperature in Celsius, condition (Clear/Rain/etc.), and humidity. We classify temperature into 5 weight zones (light to heavy) and compare each garment's fabric_weight attribute to the required weight. The suitability score drops proportionally as the fabric weight diverges from the ideal. If it's raining and the outfit has no outerwear, a penalty is applied."),
    ]

    for q,a in qas:
        h3(doc,q)
        body(doc,a)
        doc.add_paragraph()

    pb(doc)
    h2(doc,"F.4  Project Statistics Summary")
    info(doc,"These numbers show the scope and scale of the project.","E3F2FD")
    tbl(doc,["Metric","Value"],[
        ["Total Python files","34"],
        ["Total lines of code (approx.)","5,000+"],
        ["API endpoints","15+"],
        ["Database tables","4"],
        ["AI models used","2 (BlazeFace, PoseLandmarker)"],
        ["ML algorithms","KMeans clustering (sklearn)"],
        ["Scoring dimensions","5 + 1 trending bonus"],
        ["Max outfit score","105 points"],
        ["Starter wardrobe items","32 curated garments"],
        ["Fashion colour names","100+"],
        ["Python libraries","16 major libraries"],
        ["Frontend pages","8 (including landing)"],
        ["Garment categories","7"],
        ["Supported body shapes","6"],
        ["Supported skin tone levels","6"],
        ["Supported skin undertones","3"],
        ["Weekly planner cooldown","2 days"],
    ])

    h2(doc,"F.5  What Makes This Project Unique")
    bul(doc,"Multi-dimensional AI scoring — most fashion apps use one criterion (colour only or popularity only); this uses 5 simultaneous criteria")
    bul(doc,"Eastern style support — rare in fashion AI apps; full traditional garment pairing and hijab coordination")
    bul(doc,"Stochastic selection — prevents the 'always same outfit' problem using mathematical probability theory")
    bul(doc,"Real-time weather integration — garment recommendations are contextually aware of live conditions")
    bul(doc,"Body shape personalisation — uses actual pose detection landmarks, not just user self-report")
    bul(doc,"Skin tone undertone analysis — goes beyond simple dark/light to Warm/Cool/Neutral undertones for more precise colour advice")
    bul(doc,"Starter wardrobe system — new users have a working experience immediately without needing to upload anything")
    bul(doc,"Zero extra dataset needed — the scoring rules are embedded as expert knowledge in config.py, not learned from labelled data")

    pb(doc)
    h2(doc,"F.6  Potential Future Improvements")
    body(doc,"If this project were to be developed further:")
    bul(doc,"Replace rule-based body shape classification with a trained CNN that directly outputs body shape from the full image")
    bul(doc,"Add collaborative filtering — learn from thousands of users' preferences to improve colour recommendations")
    bul(doc,"Virtual try-on using image generation AI — overlay selected garments onto the user's body photo")
    bul(doc,"Mobile app using React Native frontend instead of Streamlit")
    bul(doc,"Multi-language support for UI and garment descriptions")
    bul(doc,"Shopping integration — link recommended outfit pieces to purchase links when the user doesn't own them")
    bul(doc,"Social sharing feature — share weekly outfit plans with friends")
    bul(doc,"Expand to body photo analysis using size estimation for better fit recommendations")


if __name__=="__main__":
    doc=Document()
    build_part_d(doc)
    build_part_e(doc)
    build_part_f(doc)
    out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"_master_p3.docx")
    doc.save(out)
    print(f"Part 3 saved → {out}")
